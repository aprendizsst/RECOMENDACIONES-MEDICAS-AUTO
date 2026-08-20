import streamlit as st
import streamlit.components.v1 as components
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import sqlite3
import hashlib
import os
import datetime
import re
import unicodedata
import json
import html
from pathlib import Path

try:
    import pytesseract
    from PIL import ImageEnhance, ImageFilter, ImageOps
    OCR_DISPONIBLE = True
except Exception:
    pytesseract = None
    ImageEnhance = None
    ImageFilter = None
    ImageOps = None
    OCR_DISPONIBLE = False

import requests
import io
import zipfile
import tempfile
import subprocess
import base64
import smtplib
import ssl
from email.message import EmailMessage
from email.utils import parseaddr

# ==============================================================================
# 1. CONSTANTES Y DICCIONARIOS GLOBALES (ÁMBITO SUPERIOR)
# ==============================================================================

_ETIQUETAS_CORTE = [
    "APELLIDOS Y NOMBRES", "NOMBRES Y APELLIDOS", "NOMBRE DEL TRABAJADOR",
    "NOMBRE TRABAJADOR", "NOMBRE COMPLETO", "TRABAJADOR", "PACIENTE",
    "CARGO ACTUAL", "CARGO", "OCUPACIÓN", "OCUPACION", "OFICIO", "PUESTO",
    "DOCUMENTO", "IDENTIFICACIÓN", "IDENTIFICACION", "CÉDULA", "CEDULA", "C.C.",
    "CC", "GÉNERO", "GENERO", "EDAD", "TELÉFONO", "TELEFONO", "CELULAR",
    "EPS", "AFP", "ARL", "EMPRESA", "NIT", "FECHA", "CIUDAD", "MUNICIPIO",
    "LUGAR", "SEDE", "DIRECCIÓN", "DIRECCION", "TIPO DE EXAMEN", "TIPO EXAMEN", "EVALUACION"
]

_RUIDO_IDENTIDAD = {
    "DATOS", "DATOS DEL TRABAJADOR", "INFORMACION DEL TRABAJADOR",
    "INFORMACIÓN DEL TRABAJADOR", "APELLIDOS Y NOMBRES", "NOMBRES Y APELLIDOS",
    "NOMBRE", "TRABAJADOR", "PACIENTE", "GENERO", "GÉNERO", "EDAD",
    "DOCUMENTO", "IDENTIFICACION", "IDENTIFICACIÓN", "CEDULA", "CÉDULA",
    "EMPRESA", "IPS", "EPS", "AFP", "ARL", "FIRMA", "CERTIFICADO"
}

_RUIDO_CARGO = {
    "CARGO", "CARGO ACTUAL", "OCUPACION", "OCUPACIÓN", "OFICIO", "PUESTO",
    "TRABAJADOR", "DATOS", "EMPRESA", "EPS", "AFP", "ARL", "GENERO", "GÉNERO",
    "DOCUMENTO", "IDENTIFICACION", "IDENTIFICACIÓN", "CERTIFICADO",
    "TIPO DE EXAMEN", "TIPO DE EXÁMEN", "TIPO EXAMEN", "TIPO DE EVALUACION",
    "TIPO DE EVALUACIÓN", "EVALUACION", "EVALUACIÓN", "PERIODICO", "PERIÓDICO",
    "INGRESO", "EGRESO", "RETIRO", "CAMBIO DE CARGO"
}

_RUIDO_LUGAR = {
    "LUGAR", "CIUDAD", "MUNICIPIO", "SEDE", "FECHA", "DIA", "DÍA", "MES",
    "AÑO", "ANO", "REALIZACION", "REALIZACIÓN", "EXAMEN", "EXÁMEN",
    "CERTIFICADO", "PAGINA", "PÁGINA", "LOGOTIPO", "AM", "PM", "HORA"
}

_ETIQUETAS_NOMBRE_CERTIFICADO = [
    "NOMBRES Y APELLIDOS TRABAJADOR", "NOMBRES Y APELLIDOS DEL TRABAJADOR", 
    "APELLIDOS Y NOMBRES TRABAJADOR", "APELLIDOS Y NOMBRES DEL TRABAJADOR", 
    "NOMBRES Y APELLIDOS", "APELLIDOS Y NOMBRES", "NOMBRE DEL TRABAJADOR", 
    "NOMBRE COMPLETO", "PACIENTE"
]

_ETIQUETAS_CARGO_CERTIFICADO = [
    "CARGO DEL TRABAJADOR", "CARGO ACTUAL DEL TRABAJADOR", "CARGO ACTUAL", 
    "CARGO U OCUPACIÓN", "CARGO U OCUPACION", "OCUPACIÓN DEL TRABAJADOR", 
    "OCUPACION DEL TRABAJADOR", "PUESTO DE TRABAJO", "OCUPACIÓN", "OCUPACION", 
    "OFICIO", "LABOR", "CARGO"
]

_ETIQUETAS_FECHA_CERTIFICADO = [
    "FECHA DE REALIZACIÓN DEL EXAMEN", "FECHA DE REALIZACION DEL EXAMEN", 
    "FECHA DE REALIZACIÓN DE LOS EXÁMENES", "FECHA DE REALIZACION DE LOS EXAMENES", 
    "FECHA DEL EXAMEN", "FECHA EXAMEN", "FECHA DE ATENCIÓN", "FECHA DE ATENCION"
]

_ETIQUETAS_LUGAR_CERTIFICADO = [
    "FECHA Y CIUDAD DE REALIZACIÓN", "FECHA Y CIUDAD DE REALIZACION", 
    "CIUDAD DE REALIZACIÓN DEL EXAMEN", "CIUDAD DE REALIZACION DEL EXAMEN", 
    "LUGAR DE REALIZACIÓN DEL EXAMEN", "LUGAR DE REALIZACION DEL EXAMEN", 
    "LUGAR DE REALIZACIÓN DE LOS EXÁMENES", "LUGAR DE REALIZACION DE LOS EXAMENES", 
    "LUGAR DONDE SE REALIZARON LOS EXÁMENES", "LUGAR DONDE SE REALIZARON LOS EXAMENES", 
    "MUNICIPIO DE REALIZACIÓN", "MUNICIPIO DE REALIZACION", "CIUDAD DEL EXAMEN", 
    "MUNICIPIO DEL EXAMEN", "LUGAR DEL EXAMEN", "SEDE DE ATENCIÓN", "SEDE DE ATENCION", 
    "CIUDAD", "MUNICIPIO", "LUGAR", "SEDE"
]

_ETIQUETAS_IPS_CERTIFICADO = [
    "IPS QUE REALIZA EL EXAMEN", "IPS PRESTADORA", "CENTRO MÉDICO", 
    "CENTRO MEDICO", "INSTITUCIÓN PRESTADORA", "INSTITUCION PRESTADORA"
]

_PATRONES_LEGALES_RECOMENDACIONES = [
    r"\bconsentimiento(?:\s+informado)?\b", r"\bautorizo\b", 
    r"\bautorización\s+para\s+el\s+tratamiento\s+de\s+datos\b", 
    r"\bautorizacion\s+para\s+el\s+tratamiento\s+de\s+datos\b", 
    r"\btratamiento\s+de\s+datos(?:\s+personales)?\b", r"\bprotección\s+de\s+datos\b", 
    r"\bproteccion\s+de\s+datos\b", r"\bhabeas\s+data\b", r"\bley\s+1581\b", 
    r"\bdeclaro\b", r"\bmanifiesto\b", r"\bhe\s+sido\s+informad[oa]\b", 
    r"\bacepto\s+(?:el|la|los|las)\b", r"\bconstancia\b", r"\briesgos\s+y\s+beneficios\b", 
    r"\bfirma\s+(?:del|de\s+la)\s+(?:trabajador|paciente|usuario|evaluado)\b", 
    r"\bfirma\s+del\s+m[eé]dico\b", r"\bhuella\b", r"\bdocumento\s+de\s+identidad\b", 
    r"\bresponsabilidad\s+del\s+paciente\b", r"\bdeclaración\s+del\s+paciente\b", 
    r"\bdeclaracion\s+del\s+paciente\b", r"\bderechos\s+y\s+deberes\b", 
    r"\binformación\s+suministrada\s+es\s+verdadera\b", r"\binformacion\s+suministrada\s+es\s+verdadera\b"
]

_ENCABEZADOS_LEGALES = [
    "CONSENTIMIENTO INFORMADO", "CONSENTIMIENTO", "AUTORIZACIÓN PARA TRATAMIENTO DE DATOS", 
    "AUTORIZACION PARA TRATAMIENTO DE DATOS", "TRATAMIENTO DE DATOS PERSONALES", 
    "DECLARACIÓN DEL PACIENTE", "DECLARACION DEL PACIENTE", "AUTORIZO", "CONSTANCIA", 
    "FIRMA DEL TRABAJADOR", "FIRMA DEL PACIENTE", "FIRMA DEL USUARIO", "HUELLA", "HABEAS DATA"
]

SVE_CLINICAL_KEYWORDS = {
    "VISUAL": "Conservación Visual", "AUDITIV": "Conservación Auditiva", 
    "RUIDO": "Conservación Auditiva", "OIDO": "Conservación Auditiva", "OÍDO": "Conservación Auditiva",
    "AUDIO": "Conservación Auditiva", "OSTEOMUSCULAR": "Prevención Osteomuscular (DME)",
    "POSTURAL": "Prevención Osteomuscular (DME)", "LUMBAR": "Prevención Osteomuscular (DME)",
    "ERGONOMIC": "Prevención Osteomuscular (DME)", "ESPALDA": "Prevención Osteomuscular (DME)",
    "DME": "Prevención Osteomuscular (DME)", "RESPIRATORI": "Conservación Respiratoria",
    "ESPIROMETR": "Conservación Respiratoria", "POLVO": "Conservación Respiratoria",
    "HUMO": "Conservación Respiratoria", "CARDIOVASCULAR": "Riesgo Cardiovascular"
}

EXAMS_MAP = {
    "AUDIOMETRIA DE TONOS": "Audiometría", "AUDIOMETRIA": "Audiometría",
    "ESPIROMETRIA": "Espirometría", "ESPIROMETRÍA": "Espirometría",
    "OPTOMETRIA": "Optometría", "OPTOMETRÍA": "Optometría",
    "VISIOMETRIA": "Visiometría", "VISIOMETRÍA": "Visiometría",
    "EXAMEN MEDICO OCUPACIONAL": "Examen Clínico Ocupacional", 
    "EXAMEN MEDICO": "Examen Clínico Ocupacional",
    "EXAMEN OCUPACIONAL ENFASIS OSTEOMUSCULAR": "Énfasis Osteomuscular",
    "ENFASIS OSTEOMUSCULAR": "Énfasis Osteomuscular", "ÉNFASIS OSTEOMUSCULAR": "Énfasis Osteomuscular",
    "ELECTROCARDIOGRAMA DE RITMO": "Electrocardiograma", "ELECTROCARDIOGRAMA": "Electrocardiograma", 
    "FROTIS": "Frotis", "CUADRO HEMATICO": "Cuadro Hemático", "CUADRO HEMÁTICO": "Cuadro Hemático",
    "COLESTEROL": "Colesterol", "TRIGLICERIDOS": "Triglicéridos", "PARCIAL DE ORINA": "Parcial de Orina",
    "VSH": "VSH", "PCR": "PCR",
    "GLICEMIA": "Glicemia", "GLUCOSA": "Glicemia", "HEMOGRAMA": "Hemograma",
    "HEMOGRAMA COMPLETO": "Hemograma", "PERFIL LIPIDICO": "Perfil lipídico",
    "PERFIL LIPÍDICO": "Perfil lipídico", "HDL": "Colesterol HDL",
    "LDL": "Colesterol LDL", "CREATININA": "Creatinina", "TRANSAMINASAS": "Transaminasas",
    "TGO": "TGO", "TGP": "TGP", "ACIDO URICO": "Ácido úrico", "ÁCIDO ÚRICO": "Ácido úrico",
    "BACILOSCOPIA": "Baciloscopia", "KOH": "KOH", "COPROLOGICO": "Coprológico",
    "COPROLÓGICO": "Coprológico", "TEST DE COLOR": "Test de color", "ISHIHARA": "Test de Ishihara",
    "RAYOS X": "Rayos X", "RX DE TORAX": "Rayos X de tórax", "RX DE TÓRAX": "Rayos X de tórax",
    "ELECTROENCEFALOGRAMA": "Electroencefalograma", "PSICOSENSOMETRICO": "Psicosensométrico",
    "PSICOSENSOMÉTRICO": "Psicosensométrico", "VALORACION PSICOLOGICA": "Valoración psicológica",
    "VALORACIÓN PSICOLÓGICA": "Valoración psicológica"
}

GEMINI_MODEL_DEFAULT = "gemini-3.6-flash"
GEMINI_INTERACTIONS_URL = "https://generativelanguage.googleapis.com/v1beta/interactions"
PROCESSING_PIPELINE_VERSION = "2026-08-10.7"
APP_VERSION = "2026-08-11.2"

PLANTILLA_ASUNTO_CORREO = "Recomendación médica ocupacional - {nombre}"
PLANTILLA_CUERPO_CORREO = (
    "Cordial saludo,\n\n"
    "A continuación hago envío de la recomendación médica de {nombre}, "
    "identificado(a) con el número de cédula {identificacion}.\n\n"
    "Se requiere confirmar la recepción de este correo. Asimismo, el documento debe "
    "firmarse y enviarse nuevamente de forma física, diligenciado con nombre, cédula y fecha.\n\n"
    "Cordialmente,\n"
    "Seguridad y Salud en el Trabajo\n"
    "JER S.A."
)

CAMPOS_QUE_AFECTAN_SALIDA = {
    "nombre", "cargo", "tipo_examen", "identificacion", "examenes_lista",
    "recomendaciones_por_examen", "recomendaciones_lista",
    "recomendaciones_pendientes_revision", "observaciones", "remisiones",
    "vigilancia_programa", "lugar", "fecha"
}

# --- CONFIGURACIÓN DE PÁGINA Y CSS ULTRA-RESPONSIVE ---
st.set_page_config(
    page_title="Portal SST - JER S.A.", 
    page_icon="🩺", 
    layout="wide"
)

st.markdown("""
    <style>
    html, body, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
        background-color: #0b0f19 !important;
        color: #f8fafc !important;
    }
    
    /* Ampliación de pantalla al 98% de ancho con padding equilibrado */
    .main .block-container {
        max-width: 98% !important;
        padding-top: 1.2rem !important;
        padding-bottom: 2.5rem !important;
        padding-left: 1.5rem !important;
        padding-right: 1.5rem !important;
    }
    
    /* Diseño de Columnas como Tarjetas Independientes con sombra y bordes */
    [data-testid="column"] {
        background-color: #111827 !important;
        padding: 1.8rem !important;
        border-radius: 14px !important;
        border: 1px solid #1f2937 !important;
        box-shadow: 0 10px 25px rgba(0, 0, 0, 0.3) !important;
    }
    
    .login-box {
        max-width: 480px;
        margin: 60px auto;
        padding: 40px;
        background-color: #111827 !important;
        border-radius: 16px;
        box-shadow: 0 20px 40px rgba(0, 0, 0, 0.4);
        border: 1px solid #1f2937;
    }
    
    .login-box h2 {
        color: #3b82f6 !important;
        text-align: center;
        margin-bottom: 5px;
        font-weight: 700;
        font-size: 1.8rem;
    }
    
    div[data-testid="stRadio"] label p {
        color: #f3f4f6 !important;
        font-weight: 600 !important;
        font-size: 1.1rem !important;
    }
    
    div[data-testid="stWidgetLabel"] p {
        color: #60a5fa !important;
        font-weight: 700 !important;
        font-size: 1.1rem !important;
        margin-bottom: 4px !important;
    }
    
    div[data-baseweb="input"] {
        background-color: #1f2937 !important;
        border: 1px solid #374151 !important;
        border-radius: 8px !important;
        min-height: 46px !important;
    }
    div[data-baseweb="input"] input {
        color: #ffffff !important;
        background-color: #1f2937 !important;
        font-size: 1.05rem !important;
    }
    div[data-testid="stTextArea"] textarea {
        color: #ffffff !important;
        background-color: #1f2937 !important;
        border: 1px solid #374151 !important;
        border-radius: 8px !important;
        font-size: 1.05rem !important;
        line-height: 1.5 !important;
    }
    
    button[data-baseweb="tab"] p {
        color: #9ca3af !important;
        font-size: 1.05rem !important;
    }
    button[data-baseweb="tab"][aria-selected="true"] p {
        color: #3b82f6 !important;
        font-weight: 700 !important;
    }
    
    .stButton>button {
        background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%) !important;
        color: #ffffff !important;
        border-radius: 8px !important;
        padding: 14px 28px !important;
        font-size: 1.1rem !important;
        font-weight: 700 !important;
        border: none !important;
        transition: all 0.2s ease-in-out !important;
        box-shadow: 0 4px 14px rgba(37, 99, 235, 0.3) !important;
        width: 100%;
    }
    .stButton>button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 20px rgba(37, 99, 235, 0.5) !important;
        filter: brightness(115%);
    }
    .stButton>button:active, div[data-testid="stFormSubmitButton"] button:active {
        transform: translateY(1px) scale(.99) !important;
    }
    .stButton>button:disabled, div[data-testid="stFormSubmitButton"] button:disabled {
        opacity: .48 !important;
        cursor: not-allowed !important;
        box-shadow: none !important;
        transform: none !important;
    }
    div[data-testid="stFormSubmitButton"] button {
        background: linear-gradient(135deg, #7c3aed, #4f46e5) !important;
        color: #fff !important;
        border: 0 !important;
        border-radius: 10px !important;
        min-height: 48px !important;
        font-weight: 800 !important;
        transition: transform .18s ease, box-shadow .18s ease, filter .18s ease !important;
        box-shadow: 0 6px 18px rgba(79, 70, 229, .30) !important;
    }
    div[data-testid="stFormSubmitButton"] button:hover {
        transform: translateY(-2px) !important;
        filter: brightness(112%) !important;
        box-shadow: 0 9px 24px rgba(79, 70, 229, .48) !important;
    }
    
    .header-banner {
        background: linear-gradient(135deg, #1e3a8a 0%, #0f172a 100%);
        padding: 24px;
        border-radius: 12px;
        color: #ffffff !important;
        margin-bottom: 25px;
        border: 1px solid #1e2937;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
    }
    .header-banner h1 {
        font-size: 2.3rem !important;
        margin-bottom: 6px !important;
    }
    .header-banner p {
        font-size: 1.1rem !important;
        margin: 0 !important;
    }
    
    .metric-card {
        background-color: #111827 !important;
        padding: 18px !important;
        border-radius: 10px;
        box-shadow: 0 4px 10px rgba(0,0,0,0.2);
        border-left: 5px solid #2563eb;
        margin-bottom: 12px;
        color: #e5e7eb !important;
        font-size: 1.1rem !important;
    }
    div[data-testid="stMetric"] {
        background: linear-gradient(145deg, #111827, #0f172a) !important;
        border: 1px solid #26344d !important;
        border-radius: 13px !important;
        padding: 15px 17px !important;
        box-shadow: 0 7px 18px rgba(0,0,0,.20) !important;
        transition: transform .18s ease, border-color .18s ease !important;
    }
    div[data-testid="stMetric"]:hover {
        transform: translateY(-2px) !important;
        border-color: #3b82f6 !important;
    }
    
    div[data-testid="stExpander"] {
        background-color: #111827 !important;
        border: 1px solid #1f2937 !important;
        border-radius: 8px !important;
    }
    div[data-testid="stExpander"]:hover {
        border-color: #334d78 !important;
    }
    div[data-testid="stTabs"] [data-baseweb="tab-list"] {
        gap: 7px !important;
        background: #0f172a !important;
        border: 1px solid #243047 !important;
        padding: 7px !important;
        border-radius: 13px !important;
        overflow-x: auto !important;
    }
    div[data-testid="stTabs"] button[data-baseweb="tab"] {
        border-radius: 9px !important;
        padding: 9px 14px !important;
        transition: background .18s ease, transform .18s ease !important;
    }
    div[data-testid="stTabs"] button[data-baseweb="tab"]:hover {
        background: #17233a !important;
        transform: translateY(-1px) !important;
    }
    div[data-testid="stTabs"] button[data-baseweb="tab"][aria-selected="true"] {
        background: linear-gradient(135deg, #1d4ed8, #2563eb) !important;
    }
    div[data-testid="stTabs"] button[data-baseweb="tab"][aria-selected="true"] p {
        color: #fff !important;
    }
    .email-card {
        background: linear-gradient(145deg, #111827, #0f172a);
        border: 1px solid #26344d;
        border-left: 4px solid #8b5cf6;
        border-radius: 12px;
        padding: 13px 16px;
        margin: 8px 0;
        color: #dbeafe;
    }
    .email-card strong { color:#fff; }
    .email-card small { color:#94a3b8; }
    
    [data-testid="stSidebar"] {
        min-width: 320px !important;
    }
    [data-testid="stSidebar"] > div:first-child {
        background: linear-gradient(180deg, #0f172a 0%, #111827 100%) !important;
        border-right: 1px solid #243047;
    }
    .status-row {
        display: grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap: 12px;
        margin: 0 0 22px 0;
    }
    .status-card {
        background: linear-gradient(145deg, #111827, #0f172a);
        border: 1px solid #26344d;
        border-radius: 12px;
        padding: 14px 16px;
        color: #dbeafe;
    }
    .status-card strong { display:block; color:#fff; font-size:1rem; margin-bottom:4px; }
    .status-ok { color:#4ade80; }
    .status-warn { color:#fbbf24; }
    .section-title {
        display:flex; align-items:center; gap:10px; margin:4px 0 14px;
        color:#eaf2ff; font-size:1.25rem; font-weight:800;
    }
    .section-number {
        display:inline-flex; align-items:center; justify-content:center; width:30px; height:30px;
        border-radius:9px; background:#2563eb; color:white; font-size:.95rem;
    }
    .preview-shell {
        background:#e5e7eb; border-radius:14px; padding:12px; border:1px solid #334155;
    }
    div[data-testid="stDownloadButton"] button {
        background: linear-gradient(135deg, #059669, #047857) !important;
        color:#fff !important; border:none !important; font-weight:800 !important;
    }
    @media (max-width: 900px) {
        .status-row { grid-template-columns: 1fr; }
        .main .block-container { padding-left:.8rem !important; padding-right:.8rem !important; }
    }
    </style>
""", unsafe_allow_html=True)

# --- BASE DE DATOS Y SEGURIDAD ---
DB_NAME = os.getenv("PORTAL_SST_DB", "usuarios.db")

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS usuarios 
                 (usuario TEXT PRIMARY KEY, contrasena TEXT, nombre TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS configuracion 
                 (clave TEXT PRIMARY KEY, valor TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS archivos_config
                 (clave TEXT PRIMARY KEY, nombre TEXT NOT NULL, mime TEXT,
                  contenido BLOB NOT NULL, actualizado_en TEXT NOT NULL)''')
    c.execute('''CREATE TABLE IF NOT EXISTS historial_correos
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, fecha TEXT NOT NULL,
                  pdf_origen TEXT, trabajador TEXT, destinatario TEXT,
                  archivo TEXT, estado TEXT NOT NULL, detalle TEXT,
                  cc TEXT DEFAULT '', cco TEXT DEFAULT '', asunto TEXT DEFAULT '')''')
    columnas_historial = {
        fila[1] for fila in c.execute("PRAGMA table_info(historial_correos)").fetchall()
    }
    for columna in ("cc", "cco", "asunto"):
        if columna not in columnas_historial:
            c.execute(
                f"ALTER TABLE historial_correos ADD COLUMN {columna} TEXT DEFAULT ''"
            )
    conn.commit()
    conn.close()

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def registrar_usuario(user, pwd, name):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO usuarios VALUES (?, ?, ?)", (user.lower().strip(), hash_password(pwd), name.strip()))
        conn.commit()
        success = True
    except sqlite3.IntegrityError:
        success = False
    conn.close()
    return success

def verificar_usuario(user, pwd):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT nombre FROM usuarios WHERE usuario = ? AND contrasena = ?", (user.lower().strip(), hash_password(pwd)))
    resultado = c.fetchone()
    conn.close()
    return resultado[0] if resultado else None

def actualizar_contrasena(user, old_pwd, new_pwd):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT contrasena FROM usuarios WHERE usuario = ?", (user.lower().strip(),))
    resultado = c.fetchone()
    if resultado and resultado[0] == hash_password(old_pwd):
        c.execute("UPDATE usuarios SET contrasena = ? WHERE usuario = ?", (hash_password(new_pwd), user.lower().strip()))
        conn.commit()
        conn.close()
        return True
    conn.close()
    return False

def tiene_usuarios():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM usuarios")
    count = c.fetchone()[0]
    conn.close()
    return count > 0

def guardar_config(clave, valor):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("INSERT OR REPLACE INTO configuracion VALUES (?, ?)", (clave, valor))
    conn.commit()
    conn.close()

def obtener_config(clave):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT valor FROM configuracion WHERE clave = ?", (clave,))
    res = c.fetchone()
    conn.close()
    return res[0] if res else ""

def registrar_historial_correo(
    pdf_origen, trabajador, destinatario, archivo, estado, detalle="",
    cc="", cco="", asunto=""
):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute(
        """INSERT INTO historial_correos
           (fecha, pdf_origen, trabajador, destinatario, archivo, estado, detalle, cc, cco, asunto)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (
            datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), pdf_origen,
            trabajador, destinatario, archivo, estado, str(detalle or "")[:1000],
            str(cc or "")[:500], str(cco or "")[:500], str(asunto or "")[:500]
        )
    )
    conn.commit()
    conn.close()

def obtener_historial_correos(limite=200):
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute(
        """SELECT fecha, trabajador, destinatario, cc, cco, asunto, archivo, estado, detalle
           FROM historial_correos ORDER BY id DESC LIMIT ?""",
        (int(limite),)
    )
    filas = [dict(fila) for fila in c.fetchall()]
    conn.close()
    return filas

def guardar_archivo_config(clave, archivo):
    """Guarda plantilla o firma en SQLite para conservarla entre sesiones."""
    if archivo is None:
        return False
    if isinstance(archivo, dict):
        contenido = archivo.get("bytes", b"")
        nombre = archivo.get("name", clave)
        mime = archivo.get("mime", "application/octet-stream")
    else:
        contenido = archivo.getvalue() if hasattr(archivo, "getvalue") else archivo.read()
        nombre = getattr(archivo, "name", clave)
        mime = getattr(archivo, "type", "application/octet-stream")
    if not contenido:
        return False
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute(
        """INSERT OR REPLACE INTO archivos_config
           (clave, nombre, mime, contenido, actualizado_en) VALUES (?, ?, ?, ?, ?)""",
        (clave, nombre, mime, sqlite3.Binary(contenido), datetime.datetime.now().isoformat(timespec="seconds"))
    )
    conn.commit()
    conn.close()
    return True

def obtener_archivo_config(clave):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT nombre, mime, contenido, actualizado_en FROM archivos_config WHERE clave = ?", (clave,))
    row = c.fetchone()
    conn.close()
    if not row:
        return None
    return {"name": row[0], "mime": row[1], "bytes": bytes(row[2]), "updated_at": row[3]}

def eliminar_archivo_config(clave):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("DELETE FROM archivos_config WHERE clave = ?", (clave,))
    eliminado = c.rowcount > 0
    conn.commit()
    conn.close()
    return eliminado

def archivo_a_buffer(archivo):
    """Devuelve un buffer nuevo; evita punteros agotados de UploadedFile."""
    if archivo is None:
        return None
    if isinstance(archivo, dict):
        contenido = archivo.get("bytes", b"")
        nombre = archivo.get("name", "archivo")
    elif isinstance(archivo, (bytes, bytearray)):
        contenido = bytes(archivo)
        nombre = "archivo"
    else:
        contenido = archivo.getvalue() if hasattr(archivo, "getvalue") else archivo.read()
        nombre = getattr(archivo, "name", "archivo")
    buffer = io.BytesIO(contenido)
    buffer.name = nombre
    return buffer

init_db()

# --- ESTADOS Y SESIÓN DE STREAMLIT ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "documentos" not in st.session_state:
    st.session_state.documentos = {}
if "pdfs_raw_bytes" not in st.session_state:
    st.session_state.pdfs_raw_bytes = {}
if "textos_raw" not in st.session_state:
    st.session_state.textos_raw = {}
if "export_bytes" not in st.session_state:
    st.session_state.export_bytes = None
if "zip_bytes" not in st.session_state:
    st.session_state.zip_bytes = None
if "batch_outputs" not in st.session_state:
    st.session_state.batch_outputs = {}
if "output_fingerprints" not in st.session_state:
    st.session_state.output_fingerprints = {}
if "last_cache_stats" not in st.session_state:
    st.session_state.last_cache_stats = {"generados": 0, "reutilizados": 0}
if "processed_doc" not in st.session_state:
    st.session_state.processed_doc = None
if "prev_colaborador" not in st.session_state:
    st.session_state.prev_colaborador = None
if "document_count" not in st.session_state:
    st.session_state.document_count = 0
if "ai_validation" not in st.session_state:
    st.session_state.ai_validation = {}
if "editor_version" not in st.session_state:
    st.session_state.editor_version = 0
if "asset_hashes" not in st.session_state:
    st.session_state.asset_hashes = {}
if "original_batch_preview" not in st.session_state:
    st.session_state.original_batch_preview = {}
if "email_send_feedback" not in st.session_state:
    st.session_state.email_send_feedback = None

# --- FUNCIONES DE LIMPIEZA DE TEXTO Y ORTOGRAFÍA ---
def corregir_ortografia_sst(texto):
    if not texto: return ""
    diccionario_SST = {
        r'\brealziado\b': 'realizado', r'\brealziados\b': 'realizados',
        r'\baudiometria\b': 'audiometría', r'\bvisiometria\b': 'visiometría',
        r'\bespirometria\b': 'espirometría', r'\boptometria\b': 'optometría',
        r'\bfisica\b': 'física', r'\bmedico\b': 'médico', r'\bperiodico\b': 'periódico',
        r'\bproteccion\b': 'protección', r'\balimentacion\b': 'alimentación',
        r'\brecomendacion\b': 'recomendación', r'\bperfil\s+lipidico\b': 'perfil lipídico',
        r'\benfasis\b': 'énfasis', r'\bosteomuscular\b': 'osteomuscular',
        r'\bregion\b': 'región', r'\bhabitos\b': 'hábitos', r'\badiministrativo\b': 'administrativo',
        r'\brecomendaicones\b': 'recomendaciones', r'\brecomendacines\b': 'recomendaciones',
        r'\balteracion\b': 'alteración', r'\brestriccion\b': 'restricción',
        r'\bvaloracion\b': 'valoración', r'\bevaluacion\b': 'evaluación',
        r'\bocupacionalres\b': 'ocupacionales', r'\bprotecion\b': 'protección',
        r'\bseguimineto\b': 'seguimiento', r'\bperiodicamente\b': 'periódicamente',
        r'\bconservacion\b': 'conservación', r'\boptica\b': 'óptica',
        r'\bprotecion\b': 'protección', r'\bresolucion\b': 'resolución'
    }
    for patron, reemplazo in diccionario_SST.items():
        texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)
    return texto

def a_caso_oracion(texto):
    """Normaliza mayúsculas y ortografía sin alterar el significado del contenido."""
    if not texto:
        return ""
    texto_min = corregir_ortografia_sst(str(texto)).lower()
    texto_min = re.sub(r"[ \t]+", " ", texto_min)
    texto_min = re.sub(r"\s+([,.;:!?])", r"\1", texto_min)
    texto_min = re.sub(r"([,.;:!?])(?=[^\s\n])", r"\1 ", texto_min)
    texto_min = re.sub(r"\n{3,}", "\n\n", texto_min).strip(" \n\t-_:;")
    texto_min = re.sub(
        r'(^|[:.!?]\s+|\n+)([a-zñáéíóúü])',
        lambda m: m.group(1) + m.group(2).upper(),
        texto_min
    )
    for sigla in ["SST", "PVE", "DME", "EPP", "RX", "VSH", "PCR", "TGO", "TGP", "HDL", "LDL", "IMC", "EPS", "ARL", "AFP"]:
        texto_min = re.sub(rf"\b{sigla.lower()}\b", sigla, texto_min, flags=re.IGNORECASE)
    return texto_min

def eliminar_repeticiones_internas(texto):
    """Elimina oraciones repetidas dentro del mismo elemento sin resumir su contenido."""
    texto = re.sub(r"\s+", " ", str(texto or "")).strip()
    if not texto:
        return ""
    segmentos = [s.strip() for s in re.split(r"(?<=[.!?])\s+", texto) if s.strip()]
    if len(segmentos) > 1:
        unicos = []
        firmas = set()
        for segmento in segmentos:
            firma = normalizar_etiqueta(segmento)
            if firma and firma not in firmas:
                firmas.add(firma)
                unicos.append(segmento)
        texto = " ".join(unicos)

    palabras = texto.split()
    if len(palabras) >= 8 and len(palabras) % 2 == 0:
        mitad = len(palabras) // 2
        primera = " ".join(palabras[:mitad])
        segunda = " ".join(palabras[mitad:])
        if normalizar_etiqueta(primera) == normalizar_etiqueta(segunda):
            texto = primera
    return texto

def _separar_prefijo_examen(texto):
    """Separa 'Audiometría: recomendación' sin tratar cualquier dos puntos como examen."""
    texto = str(texto or "").strip()
    if ":" not in texto:
        return "", texto
    posible_prefijo, contenido = texto.split(":", 1)
    prefijo_norm = normalizar_etiqueta(posible_prefijo)
    nombres_examen = {
        normalizar_etiqueta(valor) for valor in EXAMS_MAP.values()
    } | {
        "EXAMEN CLINICO OCUPACIONAL", "EXAMEN MEDICO OCUPACIONAL",
        "PSICOSENSOMETRICO", "ENFASIS OSTEOMUSCULAR"
    }
    es_examen = len(posible_prefijo.split()) <= 6 and any(
        prefijo_norm == nombre or prefijo_norm in nombre or nombre in prefijo_norm
        for nombre in nombres_examen
    )
    return (posible_prefijo.strip(), contenido.strip()) if es_examen else ("", texto)

def recomendacion_parece_incompleta(texto):
    """Detecta cortes OCR evidentes; no intenta completar información que no está en la fuente."""
    _, contenido = _separar_prefijo_examen(texto)
    contenido = re.sub(r"\s+", " ", contenido).strip(" •\t\r\n-_:;,.()[]")
    if not contenido:
        return True
    normal = normalizar_etiqueta(contenido)
    palabras = normal.split()
    compactas_validas = {
        "CONTROL ANUAL", "REPOSOS AUDITIVOS", "EVITAR FROTE OCULAR",
        "PAUTAS DE CUIDADO AUDITIVO", "SIN RESTRICCIONES", "SIN ALTERACIONES"
    }
    if normal in compactas_validas:
        return False
    finales_incompletos = {
        "A", "AL", "ANTE", "BAJO", "CON", "DE", "DEL", "DESDE", "DURANTE",
        "EL", "EN", "ENTRE", "HACIA", "HASTA", "LA", "LAS", "LOS", "O", "PARA",
        "POR", "QUE", "SE", "SEGUN", "SIN", "SOBRE", "SU", "SUS", "TRAS", "UN",
        "UNA", "Y"
    }
    if palabras and (palabras[-1] in finales_incompletos or len(palabras[-1]) == 1):
        return True
    if len(palabras) < 3:
        return True
    if re.search(r"(?:/{2,}|\b\d+\s*[.)]\s*$)", contenido):
        return True
    return False

def separar_recomendaciones_atomicas(items):
    """Convierte listas agrupadas en una recomendación completa por elemento."""
    atomicas = []
    pendientes = []
    inicios = (
        "realizar", "utilizar", "usar", "mantener", "evitar", "continuar", "control",
        "controles", "pautas", "reposos", "capacitación", "capacitacion", "fortalecer",
        "asistir", "seguimiento", "manipulación", "manipulacion", "alternancia"
    )
    patron_inicio = "|".join(re.escape(palabra) for palabra in inicios)

    for item in items or []:
        item = re.sub(r"\s+", " ", str(item or "")).strip()
        if not item:
            continue
        prefijo, contenido = _separar_prefijo_examen(item)
        contenido = re.sub(
            r"\b(así mismo|asimismo|de igual forma|adicionalmente),\s*",
            lambda m: m.group(1) + "§ ", contenido, flags=re.IGNORECASE
        )
        contenido = re.sub(r"(^|[,;]\s*)\d+\s*[.)-]\s*", lambda m: ("|||" if m.start() == 0 else m.group(1) + "|||"), contenido)
        contenido = re.sub(r"\s*//+\s*|,\s*-\s*|\s+[–—-]\s+|\s*[•▪◦]\s*", "|||", contenido)
        contenido = re.sub(
            rf",\s*(?=(?:{patron_inicio})\b)", "|||", contenido, flags=re.IGNORECASE
        )
        contenido = contenido.replace("§", ",")
        partes_base = [parte.strip(" ,;:-") for parte in contenido.split("|||") if parte.strip(" ,;:-")]

        partes = []
        for parte in partes_base:
            oraciones = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÜÑ])", parte)
            partes.extend(oracion.strip() for oracion in oraciones if oracion.strip())

        for parte in partes:
            parte = re.sub(r"^\d+\s*[.)-]\s*", "", parte).strip(" ,;:-")
            if not parte or es_contenido_legal_recomendacion(parte):
                continue
            texto_atomico = f"{prefijo}: {parte}" if prefijo else parte
            texto_atomico = eliminar_repeticiones_internas(a_caso_oracion(texto_atomico))
            if recomendacion_parece_incompleta(texto_atomico):
                pendientes.append(texto_atomico)
            else:
                atomicas.append(texto_atomico)

    return deduplicar_textos(atomicas), deduplicar_textos(pendientes)

def canonizar_nombre_examen(nombre, examenes_realizados=None):
    """Conserva un nombre legible y lo relaciona con el catálogo o con los exámenes detectados."""
    nombre = re.sub(r"\s+", " ", str(nombre or "")).strip(" .:-")
    if not nombre:
        return "Recomendaciones generales"
    nombre_norm = normalizar_etiqueta(nombre)
    for examen in examenes_realizados or []:
        examen_norm = normalizar_etiqueta(examen)
        if nombre_norm == examen_norm or nombre_norm in examen_norm or examen_norm in nombre_norm:
            return examen
    for clave, valor in sorted(EXAMS_MAP.items(), key=lambda item: len(item[0]), reverse=True):
        clave_norm = normalizar_etiqueta(clave)
        valor_norm = normalizar_etiqueta(valor)
        if nombre_norm in {clave_norm, valor_norm}:
            return valor
    if nombre_norm in {"GENERAL", "GENERALES", "RECOMENDACIONES GENERALES"}:
        return "Recomendaciones generales"
    return a_caso_oracion(nombre)

def agrupar_recomendaciones_por_examen(examenes_realizados, recomendaciones=None, mapa_existente=None):
    """Agrupa sin inventar: cada recomendación queda bajo su examen explícito o en Generales."""
    examenes = normalizar_lista_clinica(examenes_realizados or [])
    grupos = {examen: [] for examen in examenes}

    def agregar(examen, recomendacion):
        examen = canonizar_nombre_examen(examen, examenes)
        _, contenido = _separar_prefijo_examen(recomendacion)
        atomicas, _ = separar_recomendaciones_atomicas([contenido])
        for item in atomicas:
            _, item_limpio = _separar_prefijo_examen(item)
            item_limpio = normalizar_lista_clinica([item_limpio], cerrar_con_punto=True)
            if not item_limpio:
                continue
            firma_item = normalizar_etiqueta(item_limpio[0])
            if examen == "Recomendaciones generales" and any(
                firma_item in {normalizar_etiqueta(valor) for valor in valores}
                for nombre_grupo, valores in grupos.items()
                if nombre_grupo != "Recomendaciones generales"
            ):
                continue
            if examen != "Recomendaciones generales":
                grupos["Recomendaciones generales"] = [
                    valor for valor in grupos.get("Recomendaciones generales", [])
                    if normalizar_etiqueta(valor) != firma_item
                ]
            grupos.setdefault(examen, [])
            grupos[examen].append(item_limpio[0])

    if isinstance(mapa_existente, dict):
        for examen, elementos in mapa_existente.items():
            if isinstance(elementos, str):
                elementos = [elementos]
            for recomendacion in elementos or []:
                agregar(examen, recomendacion)
    elif isinstance(mapa_existente, list):
        for registro in mapa_existente:
            if not isinstance(registro, dict):
                continue
            examen = registro.get("examen", "Recomendaciones generales")
            elementos = registro.get("recomendaciones", [])
            if isinstance(elementos, str):
                elementos = [elementos]
            for recomendacion in elementos:
                agregar(examen, recomendacion)

    for recomendacion in recomendaciones or []:
        prefijo, contenido = _separar_prefijo_examen(recomendacion)
        agregar(prefijo or "Recomendaciones generales", contenido)

    resultado = {}
    for examen, elementos in grupos.items():
        resultado[examen] = normalizar_lista_clinica(elementos, cerrar_con_punto=True)
    if not resultado.get("Recomendaciones generales"):
        resultado.pop("Recomendaciones generales", None)
    return resultado

def aplanar_recomendaciones_por_examen(mapa):
    """Mantiene compatibilidad con la lista histórica usando el formato «Examen: recomendación»."""
    resultado = []
    for examen, recomendaciones in (mapa or {}).items():
        for recomendacion in recomendaciones or []:
            if examen == "Recomendaciones generales":
                resultado.append(recomendacion)
            else:
                resultado.append(f"{examen}: {recomendacion}")
    return normalizar_lista_clinica(deduplicar_textos(resultado), cerrar_con_punto=True)

def normalizar_lista_clinica(items, cerrar_con_punto=False):
    normalizados = []
    for item in items or []:
        texto = eliminar_repeticiones_internas(a_caso_oracion(item))
        if not texto:
            continue
        if cerrar_con_punto and texto[-1:] not in ".!?":
            texto += "."
        normalizados.append(texto)
    resultado = deduplicar_textos(normalizados)
    if cerrar_con_punto:
        resultado = [texto if texto[-1:] in ".!?" else texto + "." for texto in resultado]
    return resultado

def normalizar_datos_documento(datos):
    """Aplica el mismo formato tanto a resultados de IA como al respaldo local y la edición manual."""
    resultado = dict(datos or {})
    nombre = re.sub(r"\s+", " ", str(resultado.get("nombre", "") or "")).strip()
    cargo = re.sub(r"\s+", " ", str(resultado.get("cargo", "") or "")).strip()
    lugar = re.sub(r"\s+", " ", str(resultado.get("lugar", "") or "")).strip()
    resultado["nombre"] = nombre.title()
    resultado["cargo"] = a_caso_oracion(cargo)
    resultado["lugar"] = lugar.title()
    resultado["tipo_examen"] = a_caso_oracion(resultado.get("tipo_examen", ""))
    resultado["identificacion"] = re.sub(r"\D", "", str(resultado.get("identificacion", "") or ""))
    resultado["correo"] = str(resultado.get("correo", "") or "").strip().lower()
    resultado["examenes_lista"] = normalizar_lista_clinica(resultado.get("examenes_lista", []))
    recomendaciones, pendientes = separar_recomendaciones_atomicas(resultado.get("recomendaciones_lista", []))
    mapa_recomendaciones = agrupar_recomendaciones_por_examen(
        resultado["examenes_lista"], recomendaciones, resultado.get("recomendaciones_por_examen")
    )
    resultado["recomendaciones_por_examen"] = mapa_recomendaciones
    resultado["recomendaciones_lista"] = aplanar_recomendaciones_por_examen(mapa_recomendaciones)
    pendientes_previos = list(resultado.get("recomendaciones_pendientes_revision", []) or [])
    resultado["recomendaciones_pendientes_revision"] = deduplicar_textos(pendientes_previos + pendientes)
    resultado["observaciones"] = a_caso_oracion(resultado.get("observaciones", ""))
    resultado["remisiones"] = a_caso_oracion(resultado.get("remisiones", "No")) or "No"
    programas = re.split(r"[,;\n]+", str(resultado.get("vigilancia_programa", "") or ""))
    programas = normalizar_lista_clinica(programas)
    resultado["vigilancia_programa"] = ", ".join(programas) if programas else "Ninguno"
    return resultado

def es_vacio_o_negativo(texto):
    if not texto: return True
    return texto.strip().lower().strip(" .-_/ '\"") in ["no", "ninguna", "ninguno", "no registra", "sin remisiones", "normal", "n/a", "sin remisión"]

def es_vacio_o_estado(texto):
    if not texto: return True
    t_clean = texto.strip().upper()
    t_clean_norm = re.sub(r'[^A-ZÁÉÍÓÚÑ\s]', '', t_clean).strip()
    t_clean_norm = re.sub(r'\s+', ' ', t_clean_norm)
    
    frases_estado = {
        "REALIZADO", "REALZIADO", "SIN ALTERACIONES", "NORMAL", "SANO", "NEGATIVO", 
        "NO REGISTRA", "NA", "SIN REMISIONES", "SIN REMISIÓN", "VISUAL", "CARDIOVASCULAR", 
        "DME", "OSTEOMUSCULAR", "AUDITIVO", "RESPIRATORIO", "SVE", "SISTEMA", "VIGILANCIA",
        "SANO Y SIN ALTERACIONES", "NINGUNO", "NINGUNA", "NO PRESENTAS", "NO PRESENTA", 
        "NO REGISTRA RECOMENDACIONES", "NORMALES", "NORMAL", "SIN ALTERACION", "NO APLICA",
        "RECOMENDACIONES MÉDICAS", "RECOMENDACIONES OCUPACIONALES", "HABITOS Y ESTILO DE VIDA SALUDABLES",
        "HABITOS SALUDABLES", "OTRAS OBSERVACIONES Y RECOMENDACIONES", "RECOMENDACIONES MEDICAS"
    }
    
    if t_clean_norm in frases_estado: return True
    if len(texto.strip()) <= 3: return True
    return False

def limpiar_campo(texto):
    if not texto: return ""
    partes = re.split(r'\b(Teléfono|Telefono|Tel|C\.C|CC|Documento|Cedula|Cargo|Fecha)\b', texto, flags=re.IGNORECASE)
    return re.sub(r'[:\-,_]+', '', partes[0]).strip().title()

def limpiar_linea_ruido_lateral(linea):
    patron_ruido = r'\s{2,}(VISUAL|DME|CARDIOVASCULAR|SVE|AUDITIVO|RESPIRATORIO|SISTEMA|VIGILANCIA)\s*$'
    linea_limpia = re.sub(patron_ruido, '', linea, flags=re.IGNORECASE)
    return linea_limpia.strip()

def limpiar_ruido_columnas_final(texto):
    if not texto: return ""
    patrones_ruido = [
        r'\bvisual\b', r'\bdme\b', r'\bcardiovascular\b', r'\bsve\b', 
        r'\bauditivo\b', r'\brespiratorio\b', r'\bsistema\b', r'\bvigilancia\b'
    ]
    for patron in patrones_ruido:
        texto = re.sub(patron + r'\s*$', '', texto, flags=re.IGNORECASE)
    return texto.strip(" :-,_/")

def intentar_parsear_fecha(fecha_str):
    if not fecha_str: return datetime.date.today()
    fecha_str = fecha_str.lower().strip(" :-,_/.()[]|")
    
    meses_dict = {
        "enero": 1, "ene": 1, "febrero": 2, "feb": 2, "marzo": 3, "mar": 3,
        "abril": 4, "abr": 4, "mayo": 5, "may": 5, "junio": 6, "jun": 6,
        "julio": 7, "jul": 7, "agosto": 8, "ago": 8, "septiembre": 9, "sep": 9, "sept": 9,
        "octubre": 10, "oct": 10, "noviembre": 11, "nov": 11, "diciembre": 12, "dic": 12
    }

    m_letras = re.search(r'(\d{1,2})\s+(?:de\s+)?([a-z]{3,10})\.?\s+(?:de\s+)?(20\d{2})', fecha_str)
    if m_letras:
        dia = int(m_letras.group(1))
        mes_str = m_letras.group(2).lower()
        anio = int(m_letras.group(3))
        if mes_str in meses_dict:
            try: return datetime.date(anio, meses_dict[mes_str], dia)
            except ValueError: pass

    m_ymd = re.search(r'\b(20\d{2})[-/.s](\d{1,2})[-/.s](\d{1,2})\b', fecha_str)
    if m_ymd: 
        try: return datetime.date(int(m_ymd.group(1)), int(m_ymd.group(2)), int(m_ymd.group(3)))
        except ValueError: pass

    m_dmy = re.search(r'\b(\d{1,2})[-/.s](\d{1,2})[-/.s](20\d{2})\b', fecha_str)
    if m_dmy: 
        try: return datetime.date(int(m_dmy.group(3)), int(m_dmy.group(2)), int(m_dmy.group(1)))
        except ValueError: pass

    return datetime.date.today()

# --- FUNCIONES DE FILTRADO DE METADATOS Y CANDIDATOS ---
def normalizar_etiqueta(texto):
    if not texto: return ""
    texto = unicodedata.normalize("NFKD", str(texto))
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = texto.upper()
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip(" :|_-./")

def dividir_columnas_estructuradas(linea):
    if not linea: return []
    columnas = [re.sub(r"\s+", " ", c).strip(" |/-,_.:") for c in re.split(r"\s{2,}|\t+|\|", linea)]
    return [c for c in columnas if c]

def recortar_en_siguiente_etiqueta(valor):
    if not valor: return ""
    patron = r"\b(?:" + "|".join(sorted((re.escape(e) for e in _ETIQUETAS_CORTE), key=len, reverse=True)) + r")\b"
    coincidencias = list(re.finditer(patron, valor, flags=re.IGNORECASE))
    if coincidencias:
        primera = coincidencias[0]
        if primera.start() > 0: valor = valor[:primera.start()]
    return valor.strip(" |/-,_.:")

def limpiar_candidato_campo(valor, tipo):
    if not valor: return ""
    valor = str(valor).replace("\x00", " ")
    valor = re.sub(r"\s+", " ", valor).strip(" |/-,_.:")
    valor = recortar_en_siguiente_etiqueta(valor)

    valor = re.split(
        r"\b(?:C\.?\s*C\.?|CÉDULA|CEDULA|DOCUMENTO|IDENTIFICACIÓN|IDENTIFICACION|TELÉFONO|TELEFONO|CELULAR|EDAD|GÉNERO|GENERO|EPS|AFP|ARL)\b",
        valor, maxsplit=1, flags=re.IGNORECASE
    )[0].strip(" |/-,_.:")

    if tipo in {"nombre", "cargo"}:
        valor = re.sub(r"\([^)]*\)", "", valor).strip()
    elif tipo == "lugar":
        valor = re.sub(r"\s+\([^)]*(?:PÁGINA|PAGINA|HORA|AM|PM)[^)]*\)", "", valor, flags=re.IGNORECASE)
        valor = re.sub(r"^\s*(?:(?:20\d{2})\s*[|/\-.]\s*\d{1,2}\s*[|/\-.]\s*\d{1,2}|\d{1,2}\s*[|/\-.]\s*\d{1,2}\s*[|/\-.]\s*20\d{2})\s*[|/\-,_.:]*\s*", "", valor)
        if "|" in valor:
            partes_lugar = [p.strip(" |/-,_.:") for p in valor.split("|") if p.strip(" |/-,_.:")]
            partes_con_letras = [p for p in partes_lugar if sum(c.isalpha() for c in p) >= 3]
            if partes_con_letras: valor = partes_con_letras[-1]
    return re.sub(r"\s+", " ", valor).strip(" |/-,_.:")

def candidato_nombre_valido(valor):
    valor = limpiar_candidato_campo(valor, "nombre")
    if not valor or len(valor) < 5 or len(valor) > 100: return False
    if re.search(r"\d|@|https?://", valor): return False
    tokens = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ'-]+", valor)
    if len(tokens) < 2 or len(tokens) > 8: return False
    norm = normalizar_etiqueta(valor)
    if norm in {normalizar_etiqueta(x) for x in _RUIDO_IDENTIDAD}: return False
    if any(ruido in norm for ruido in ["CERTIFICADO", "MEDICINA", "OCUPACIONAL", "EMPRESA", "INSTITUCION", "LABORATORIO", "FIRMA DEL", "NOMBRE DEL"]): return False
    return (sum(c.isalpha() for c in valor) / max(len(valor), 1)) >= 0.65

def candidato_cargo_valido(valor):
    valor = limpiar_candidato_campo(valor, "cargo")
    if not valor or len(valor) < 3 or len(valor) > 120: return False
    norm = normalizar_etiqueta(valor)
    if norm in {normalizar_etiqueta(x) for x in _RUIDO_CARGO}: return False
    if any(ruido in norm for ruido in ["CERTIFICADO", "MÉDICO", "MEDICO", "FIRMA", "DOCUMENTO", "IDENTIFICACION", "GENERO", "TIPO DE EXAMEN", "TIPO DE EXÁMEN", "EVALUACION", "EVALUACIÓN", "PERIODICO", "PERIÓDICO"]): return False
    if re.fullmatch(r"[\d\s./-]+", valor): return False
    return sum(c.isalpha() for c in valor) >= 3

def candidato_lugar_valido(valor):
    valor = limpiar_candidato_campo(valor, "lugar")
    if not valor or len(valor) < 3 or len(valor) > 100: return False
    if re.search(r"https?://|www\.|@", valor, flags=re.IGNORECASE): return False
    if re.fullmatch(r"[\d\s:./-]+", valor): return False
    norm = normalizar_etiqueta(valor)
    if norm in {normalizar_etiqueta(x) for x in _RUIDO_LUGAR}: return False
    if any(ruido in norm for ruido in ["PAGINA", "PÁGINA", "CERTIFICADO", "LOGOTIPO", "FIRMA", "HORA", "AM", "PM"]): return False
    return sum(c.isalpha() for c in valor) >= 3

def elegir_mejor_candidato(candidatos, tipo):
    validadores = {"nombre": candidato_nombre_valido, "cargo": candidato_cargo_valido, "lugar": candidato_lugar_valido}
    validador = validadores[tipo]
    mejores = []

    for puntaje, valor, origen in candidatos:
        limpio = limpiar_candidato_campo(valor, tipo)
        if not validador(limpio): continue
        tokens = limpio.split()
        if tipo == "nombre":
            if 2 <= len(tokens) <= 5: puntaje += 8
            if limpio.upper() == limpio: puntaje += 3
        elif tipo == "cargo" and 1 <= len(tokens) <= 7: puntaje += 4
        elif tipo == "lugar" and 1 <= len(tokens) <= 5: puntaje += 4
        mejores.append((puntaje, -len(limpio), limpio, origen))

    if not mejores: return ""
    mejores.sort(reverse=True)
    valor = mejores[0][2]
    if tipo == "nombre": return valor.title()
    if tipo == "cargo": return corregir_ortografia_sst(valor).title()
    return valor.title()

def _coincide_etiqueta(columna, etiquetas):
    norm = normalizar_etiqueta(columna)
    return any(norm == normalizar_etiqueta(e) or normalizar_etiqueta(e) in norm for e in etiquetas)

def extraer_campo_por_etiquetas(lineas, etiquetas, tipo):
    candidatos = []
    etiquetas_ordenadas = sorted(etiquetas, key=len, reverse=True)
    patron_etiquetas = "|".join(re.escape(e) for e in etiquetas_ordenadas)

    for idx, linea in enumerate(lineas):
        if not linea or not linea.strip(): continue

        m_inline = re.search(rf"(?:{patron_etiquetas})(?:\s*[:=]\s*|\s+-\s+|\s{{2,}}|\|)\s*(.+)$", linea, flags=re.IGNORECASE)
        if m_inline: candidatos.append((115, m_inline.group(1), "etiqueta en línea"))

        columnas_header = dividir_columnas_estructuradas(linea)
        indices = [pos for pos, col in enumerate(columnas_header) if _coincide_etiqueta(col, etiquetas)]

        if indices:
            for offset in range(1, 5):
                if idx + offset >= len(lineas): break
                siguiente = lineas[idx + offset].strip()
                if not siguiente: continue
                columnas_valor = dividir_columnas_estructuradas(siguiente)
                for pos in indices:
                    if pos < len(columnas_valor):
                        candidatos.append((140 - offset, columnas_valor[pos], "tabla encabezado/valor"))
                
                norm_linea = normalizar_etiqueta(linea)
                if any(norm_linea == normalizar_etiqueta(e) for e in etiquetas):
                    candidatos.append((100 - offset, siguiente, "línea siguiente"))
                break

        norm_linea = normalizar_etiqueta(linea)
        for etiqueta in etiquetas_ordenadas:
            norm_etiqueta = normalizar_etiqueta(etiqueta)
            if norm_linea.startswith(norm_etiqueta) and len(norm_linea) > len(norm_etiqueta):
                resto = linea[len(etiqueta):].strip(" |/-,_.:")
                if resto: candidatos.append((92, resto, "etiqueta inicial"))
                break

    return elegir_mejor_candidato(candidatos, tipo)

def extraer_fecha_y_lugar_robusto(lineas, texto_completo):
    candidatos_lugar = []
    fechas = []

    patron_fecha_dmy = re.compile(r"\b(\d{1,2})\s*[\s|/\-.]+\s*(\d{1,2})\s*[\s|/\-.]+\s*(20\d{2})\b")
    patron_fecha_ymd = re.compile(r"\b(20\d{2})\s*[\s|/\-.]+\s*(\d{1,2})\s*[\s|/\-.]+\s*(\d{1,2})\b")

    for line in lineas:
        if "NACIMIENTO" in line.upper() or "NACIDA" in line.upper(): continue

        parts = [p.strip() for p in re.split(r'\|', line) if p.strip()]
        for i in range(len(parts) - 2):
            if re.match(r'^\d{1,2}$', parts[i]) and re.match(r'^\d{1,2}$', parts[i+1]) and re.match(r'^20\d{2}$', parts[i+2]):
                try:
                    d, m, y = int(parts[i]), int(parts[i+1]), int(parts[i+2])
                    fechas.append((150, datetime.date(y, m, d)))
                    for j in range(i+3, len(parts)):
                        text_cand_clean = re.sub(r'\(.*?\)', '', parts[j]).strip(" |/-,_.:")
                        text_cand_clean = re.sub(r'\b(CIUDAD|MUNICIPIO|FECHA|REALIZACI[ÓO]N)\b', '', text_cand_clean, flags=re.IGNORECASE).strip(" |/-,_.:")
                        if text_cand_clean and len(text_cand_clean) > 2 and not any(w in text_cand_clean.upper() for w in ["PÁGINA", "PAGINA", "CERTIFICADO", "A.M.", "P.M.", "HORA"]):
                            candidatos_lugar.append((150, text_cand_clean, "grid matricial celda"))
                except ValueError: pass

    for idx, linea in enumerate(lineas):
        if "NACIMIENTO" in linea.upper(): continue
        columnas = dividir_columnas_estructuradas(linea)
        columnas_norm = [normalizar_etiqueta(c) for c in columnas]
        ciudad_indices = [i for i, c in enumerate(columnas_norm) if any(k in c for k in ["CIUDAD", "MUNICIPIO", "LUGAR", "SEDE"])]

        if ciudad_indices and any(k in " ".join(columnas_norm) for k in ["DIA", "MES", "ANO", "FECHA", "REALIZACION"]):
            for offset in range(1, 4):
                if idx + offset >= len(lineas): break
                valores = dividir_columnas_estructuradas(lineas[idx + offset])
                for pos in ciudad_indices:
                    if pos < len(valores):
                        candidatos_lugar.append((125 - offset, valores[pos], "tabla fecha/lugar"))
                linea_valores = lineas[idx + offset]
                m_dmy, m_ymd = patron_fecha_dmy.search(linea_valores), patron_fecha_ymd.search(linea_valores)
                try:
                    if m_ymd: fechas.append((120 - offset, datetime.date(int(m_ymd.group(1)), int(m_ymd.group(2)), int(m_ymd.group(3)))))
                    elif m_dmy: fechas.append((120 - offset, datetime.date(int(m_dmy.group(3)), int(m_dmy.group(2)), int(m_dmy.group(1)))))
                except ValueError: pass
                if valores: break

        m_dmy, m_ymd = patron_fecha_dmy.search(linea), patron_fecha_ymd.search(linea)
        m_fecha = m_ymd or m_dmy
        if m_fecha:
            try:
                if m_ymd: fecha_detectada = datetime.date(int(m_ymd.group(1)), int(m_ymd.group(2)), int(m_ymd.group(3)))
                else: fecha_detectada = datetime.date(int(m_dmy.group(3)), int(m_dmy.group(2)), int(m_dmy.group(1)))
                fechas.append((90, fecha_detectada))
            except ValueError: pass

            antes, despues = linea[:m_fecha.start()].strip(" |/-,_.:"), linea[m_fecha.end():].strip(" |/-,_.:")
            contexto = normalizar_etiqueta(linea)
            puntaje = 118 if any(k in contexto for k in ["REALIZACION", "CIUDAD", "MUNICIPIO", "LUGAR", "SEDE"]) else 78
            if despues and not any(m in despues.upper() for m in ["PÁGINA", "PAGINA", "CERTIFICADO", "P.M.", "A.M."]):
                candidatos_lugar.append((puntaje, despues, "después de fecha"))
            if antes and len(antes.split()) <= 7:
                antes = re.sub(r"(?i)\b(FECHA|CIUDAD|MUNICIPIO|LUGAR|REALIZACI[ÓO]N|DEL EXAMEN|DEL EXÁMEN)\b", "", antes).strip(" |/-,_.:")
                if antes: candidatos_lugar.append((puntaje - 5, antes, "antes de fecha"))

    meses = "enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre"
    patron_ciudad_fecha = re.compile(rf"\b([A-Za-zÁÉÍÓÚÜÑáéíóúüñ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ .'-]{{2,45}}),?\s+(\d{{1,2}}\s+de\s+(?:{meses})\s+de\s+20\d{{2}})\b", flags=re.IGNORECASE)
    for m in patron_ciudad_fecha.finditer(texto_completo):
        lugar = re.split(r"\n|[:;]", m.group(1).strip())[-1].strip()
        candidatos_lugar.append((112, lugar, "ciudad y fecha en letras"))
        try: fechas.append((112, intentar_parsear_fecha(m.group(2))))
        except: pass

    etiquetas_lugar = ["FECHA Y CIUDAD DE REALIZACIÓN", "FECHA Y CIUDAD DE REALIZACION", "CIUDAD DE REALIZACIÓN", "CIUDAD", "MUNICIPIO", "SEDE"]
    lugar_etiquetado = extraer_campo_por_etiquetas(lineas, etiquetas_lugar, "lugar")
    if lugar_etiquetado: candidatos_lugar.append((130, lugar_etiquetado, "etiqueta explícita"))

    lugar_final = elegir_mejor_candidato(candidatos_lugar, "lugar")
    fecha_final = max(fechas, key=lambda item: item[0])[1] if fechas else datetime.date.today()
    return fecha_final, lugar_final

def extraer_identidad_cargo_lugar(texto):
    lineas = [line.rstrip() for line in texto.splitlines()]
    etiquetas_nombre = ["APELLIDOS Y NOMBRES DEL TRABAJADOR", "NOMBRES Y APELLIDOS", "APELLIDOS Y NOMBRES", "NOMBRE DEL TRABAJADOR", "TRABAJADOR", "PACIENTE"]
    etiquetas_cargo = ["CARGO ACTUAL DEL TRABAJADOR", "CARGO DEL TRABAJADOR", "CARGO ACTUAL", "OCUPACIÓN", "OCUPACION", "PUESTO", "CARGO"]

    nombre = extraer_campo_por_etiquetas(lineas, etiquetas_nombre, "nombre")
    cargo = extraer_campo_por_etiquetas(lineas, etiquetas_cargo, "cargo")
    fecha, lugar = extraer_fecha_y_lugar_robusto(lineas, texto)

    return {"nombre": nombre, "cargo": cargo, "fecha": fecha, "lugar": lugar}

def es_contenido_legal_recomendacion(texto):
    if not texto: return False
    limpio = re.sub(r"\s+", " ", str(texto)).strip()
    return any(re.search(patron, limpio, flags=re.IGNORECASE) for patron in _PATRONES_LEGALES_RECOMENDACIONES)

def es_encabezado_legal(texto):
    if not texto: return False
    normalizado = normalizar_etiqueta(texto)
    return any(encabezado in normalizado for encabezado in _ENCABEZADOS_LEGALES)

def recortar_contenido_legal(texto):
    if not texto: return ""
    texto = str(texto)
    posiciones = []
    for patron in _PATRONES_LEGALES_RECOMENDACIONES:
        coincidencia = re.search(patron, texto, flags=re.IGNORECASE)
        if coincidencia: posiciones.append(coincidencia.start())
    if posiciones: texto = texto[:min(posiciones)]
    lineas_validas = []
    for linea in texto.splitlines():
        if es_encabezado_legal(linea) or es_contenido_legal_recomendacion(linea): break
        lineas_validas.append(linea)
    return re.sub(r"\s+", " ", " ".join(lineas_validas)).strip(" .;:-_/|")

def filtrar_recomendaciones_clinicas(recomendaciones):
    resultado = []
    vistos = set()
    for recomendacion in recomendaciones or []:
        limpia = recortar_contenido_legal(recomendacion)
        limpia = re.sub(r"^(?:Audiometría|Espirometría|Optometría|Visiometría|Examen Clínico Ocupacional|Énfasis Osteomuscular|Electrocardiograma|Frotis|Cuadro Hemático|Colesterol|Triglicéridos|Parcial de Orina|VSH|PCR)\s*:\s*$", "", limpia, flags=re.IGNORECASE).strip()
        if not limpia or es_vacio_o_estado(limpia) or es_contenido_legal_recomendacion(limpia): continue
        clave = normalizar_etiqueta(limpia)
        if clave not in vistos:
            vistos.add(clave)
            resultado.append(limpia)
    finales = []
    for idx, item in enumerate(resultado):
        norm_item = normalizar_etiqueta(item)
        contenido_en_otro = any(
            idx != otro_idx and norm_item in normalizar_etiqueta(otro) and len(otro) > len(item) + 8
            for otro_idx, otro in enumerate(resultado)
        )
        if not contenido_en_otro:
            finales.append(item)
    return finales

def _limpiar_celda_certificado(valor):
    if valor is None: return ""
    valor = str(valor).replace("\x00", " ")
    valor = re.sub(r"[ \t]+", " ", valor)
    valor = re.sub(r"\s*\n\s*", "\n", valor)
    return valor.strip(" |/-,_.:")

def _normalizar_lista_celdas(fila):
    return [_limpiar_celda_certificado(celda) for celda in (fila or [])]

def _contiene_alguna_etiqueta(texto, etiquetas):
    normalizado = normalizar_etiqueta(texto)
    return any(normalizar_etiqueta(etiqueta) in normalizado for etiqueta in etiquetas)

def _es_rotulo_general(texto):
    normalizado = normalizar_etiqueta(texto)
    rotulos = _ETIQUETAS_NOMBRE_CERTIFICADO + _ETIQUETAS_CARGO_CERTIFICADO + _ETIQUETAS_FECHA_CERTIFICADO + _ETIQUETAS_LUGAR_CERTIFICADO + _ETIQUETAS_IPS_CERTIFICADO + ["DOCUMENTO", "IDENTIFICACIÓN", "IDENTIFICACION", "CÉDULA", "CEDULA", "EDAD", "GÉNERO", "GENERO", "SEXO", "EMPRESA", "EPS", "ARL", "AFP", "DÍA", "DIA", "MES", "AÑO", "ANO", "TIPO DE EXAMEN", "TIPO DE EXÁMEN", "TIPO DE EVALUACIÓN"]
    return any(normalizado == normalizar_etiqueta(rotulo) or normalizado.startswith(normalizar_etiqueta(rotulo) + " ") for rotulo in rotulos)

def _extraer_valor_inline(celda, etiquetas):
    celda = _limpiar_celda_certificado(celda)
    if not celda: return ""
    plano = re.sub(r"\s*\n\s*", " | ", celda)
    for etiqueta in sorted(etiquetas, key=len, reverse=True):
        patron = re.compile(rf"^\s*{re.escape(etiqueta)}\s*(?:[:=\-|]\s*)?(.+)$", flags=re.IGNORECASE)
        coincidencia = patron.match(plano)
        if coincidencia:
            resto = coincidencia.group(1).strip(" |/-,_.:")
            if resto and not _es_rotulo_general(resto): return resto
        lineas = [linea.strip(" |/-,_.:") for linea in celda.splitlines() if linea.strip(" |/-,_.:")]
        if lineas and normalizar_etiqueta(lineas[0]) == normalizar_etiqueta(etiqueta) and len(lineas) > 1:
            resto = " ".join(lineas[1:]).strip(" |/-,_.:")
            if resto and not _es_rotulo_general(resto): return resto
    return ""

def _nombre_muy_valido(valor):
    limpio = limpiar_candidato_campo(valor, "nombre")
    if not candidato_nombre_valido(limpio): return False
    norm = normalizar_etiqueta(limpio)
    prohibidas = ["FECHA", "REALIZACION", "EXAMEN", "CARGO", "EMPRESA", "IDENTIFICACION", "DOCUMENTO", "CERTIFICADO", "CONCEPTO", "RECOMENDACION", "VIGILANCIA", "CONSENTIMIENTO", "TRABAJADOR", "APELLIDOS Y NOMBRES", "NOMBRES Y APELLIDOS"]
    if any(palabra in norm for palabra in prohibidas): return False
    tokens = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ'-]+", limpio)
    return 2 <= len(tokens) <= 7

def _cargo_muy_valido(valor):
    limpio = limpiar_candidato_campo(valor, "cargo")
    if not candidato_cargo_valido(limpio): return False
    norm = normalizar_etiqueta(limpio)
    prohibidas = ["FECHA", "REALIZACION", "EXAMEN", "DOCUMENTO", "IDENTIFICACION", "NOMBRES Y APELLIDOS", "APELLIDOS Y NOMBRES", "CERTIFICADO", "CONSENTIMIENTO", "DIA MES ANO", "TIPO DE EXAMEN", "EVALUACION", "PERIODICO", "PERIÓDICO"]
    return not any(palabra in norm for palabra in prohibidas)

def _lugar_muy_valido(valor):
    limpio = limpiar_candidato_campo(valor, "lugar")
    if not candidato_lugar_valido(limpio): return False
    norm = normalizar_etiqueta(limpio)
    prohibidas = ["FECHA", "REALIZACION", "EXAMEN", "DIA", "MES", "ANO", "DOCUMENTO", "IDENTIFICACION", "CONSENTIMIENTO"]
    return not any(palabra == norm for palabra in prohibidas)

def _agregar_candidato_especial(candidatos, tipo, puntaje, valor, origen):
    if not valor: return
    validadores = {"nombre": _nombre_muy_valido, "cargo": _cargo_muy_valido, "lugar": _lugar_muy_valido}
    limpio = limpiar_candidato_campo(valor, tipo)
    if validadores[tipo](limpio): candidatos[tipo].append((puntaje, limpio, origen))

def _buscar_valor_debajo(filas, fila_inicio, columna, tipo, max_filas=5):
    validadores = {"nombre": _nombre_muy_valido, "cargo": _cargo_muy_valido, "lugar": _lugar_muy_valido}
    validador = validadores[tipo]

    for salto in range(1, max_filas + 1):
        indice = fila_inicio + salto
        if indice >= len(filas): break
        fila = filas[indice]
        if not any(fila): continue

        for celda in fila:
            if not celda: continue
            limpio = limpiar_candidato_campo(celda, tipo)
            if validador(limpio): return limpio, salto
    return "", 0

def _buscar_valor_derecha(fila, columna, tipo):
    validadores = {"nombre": _nombre_muy_valido, "cargo": _cargo_muy_valido, "lugar": _lugar_muy_valido}
    validador = validadores[tipo]

    for pos in range(columna + 1, min(len(fila), columna + 6)):
        celda = fila[pos]
        if not celda or _es_rotulo_general(celda): continue
        limpio = limpiar_candidato_campo(celda, tipo)
        if validador(limpio): return limpio, pos - columna
    return "", 0

def _fecha_desde_componentes(dia, mes, anio):
    try:
        dia, mes, anio = int(str(dia).strip()), int(str(mes).strip()), int(str(anio).strip())
        if anio < 100: anio += 2000
        return datetime.date(anio, mes, dia)
    except (TypeError, ValueError): return None

def _buscar_fecha_en_texto(texto):
    if not texto: return None
    if "NACIMIENTO" in str(texto).upper(): return None
    texto = re.sub(r"\s+", " ", str(texto))

    patrones = [
        (re.compile(r"\b(20\d{2})\s*[\s|/\-.]+\s*(\d{1,2})\s*[\s|/\-.]+\s*(\d{1,2})\b"), lambda m: _fecha_desde_componentes(m.group(3), m.group(2), m.group(1))),
        (re.compile(r"\b(\d{1,2})\s*[\s|/\-.]+\s*(\d{1,2})\s*[\s|/\-.]+\s*(20\d{2})\b"), lambda m: _fecha_desde_componentes(m.group(1), m.group(2), m.group(3)))
    ]

    for patron, constructor in patrones:
        coincidencia = patron.search(texto)
        if coincidencia:
            fecha = constructor(coincidencia)
            if fecha: return fecha

    meses = {"enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6, "julio": 7, "agosto": 8, "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12}
    coincidencia = re.search(r"\b(\d{1,2})\s+de\s+([a-záéíóúüñ]+)\s+de\s+(20\d{2})\b", texto, flags=re.IGNORECASE)
    if coincidencia:
        mes = meses.get(coincidencia.group(2).lower())
        if mes: return _fecha_desde_componentes(coincidencia.group(1), mes, coincidencia.group(3))
    return None

def _buscar_fecha_tabla(filas, fila_inicio, columna_inicio=0):
    limite = min(len(filas), fila_inicio + 7)
    for indice in range(fila_inicio, limite):
        fila = filas[indice]
        if "NACIMIENTO" in " ".join(celda for celda in fila if celda).upper(): continue
        texto_fila = " | ".join(celda for celda in fila if celda)
        fecha = _buscar_fecha_en_texto(texto_fila)
        if fecha: return fecha, 250 - (indice - fila_inicio)

    for indice in range(fila_inicio, limite):
        fila = filas[indice]
        normalizadas = [normalizar_etiqueta(celda) for celda in fila]
        indice_dia = next((pos for pos, valor in enumerate(normalizadas) if valor in {"DIA", "DÍA"}), None)
        indice_mes = next((pos for pos, valor in enumerate(normalizadas) if valor == "MES"), None)
        indice_anio = next((pos for pos, valor in enumerate(normalizadas) if valor in {"ANO", "AÑO"}), None)

        if None not in (indice_dia, indice_mes, indice_anio):
            for salto in range(1, 4):
                fila_valor_idx = indice + salto
                if fila_valor_idx >= len(filas): break
                valores = filas[fila_valor_idx]
                if max(indice_dia, indice_mes, indice_anio) >= len(valores): continue
                fecha = _fecha_desde_componentes(valores[indice_dia], valores[indice_mes], valores[indice_anio])
                if fecha: return fecha, 270 - salto
    return None, 0

def _quitar_fecha_para_lugar(texto):
    if not texto: return ""
    valor = str(texto)
    valor = re.sub(r"\b20\d{2}\s*[-/.\s]\s*\d{1,2}\s*[-/.\s]\s*\d{1,2}\b", " ", valor)
    valor = re.sub(r"\b\d{1,2}\s*[-/.\s]\s*\d{1,2}\s*[-/.\s]\s*20\d{2}\b", " ", valor)
    valor = re.sub(r"\b\d{1,2}\s+de\s+[a-záéíóúüñ]+\s+de\s+20\d{2}\b", " ", valor, flags=re.IGNORECASE)
    valor = re.sub(r"\b(FECHA|DÍA|DIA|MES|AÑO|ANO|CIUDAD|MUNICIPIO|LUGAR|REALIZACIÓN|REALIZACION|DEL EXAMEN|DE LOS EXÁMENES|DE LOS EXAMENES|SEDE)\b", " ", valor, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", valor).strip(" |/-,_.:")

def _extraer_lineas_por_coordenadas(page):
    try: palabras = page.extract_words(x_tolerance=2, y_tolerance=3, keep_blank_chars=False, use_text_flow=False) or []
    except: return []

    grupos = []
    for palabra in sorted(palabras, key=lambda item: (round(float(item["top"]), 1), float(item["x0"]))):
        top = float(palabra["top"])
        grupo = None
        for existente in reversed(grupos[-10:]):
            if abs(existente["top"] - top) <= 3.2:
                grupo = existente
                break
        if grupo is None:
            grupo = {"top": top, "words": []}
            grupos.append(grupo)
        grupo["words"].append(palabra)

    lineas = []
    for grupo in grupos:
        palabras_linea = sorted(grupo["words"], key=lambda item: float(item["x0"]))
        partes = []
        x1_anterior = None
        for palabra in palabras_linea:
            if x1_anterior is not None:
                separacion = float(palabra["x0"]) - x1_anterior
                partes.append("    " if separacion > 16 else " ")
            partes.append(str(palabra["text"]))
            x1_anterior = float(palabra["x1"])
        linea = "".join(partes).strip()
        if linea: lineas.append(linea)
    return lineas

def _ocr_pagina_si_disponible(page):
    if not OCR_DISPONIBLE: return ""
    try:
        imagen = page.to_image(resolution=230).original.convert("L")
        imagen = ImageOps.autocontrast(imagen)
        imagen = ImageEnhance.Contrast(imagen).enhance(1.35)
        imagen = imagen.filter(ImageFilter.SHARPEN)
        texto = pytesseract.image_to_string(imagen, lang="spa+eng", config="--oem 1 --psm 6", timeout=80)
        return re.sub(r"\n{3,}", "\n\n", texto or "").strip()
    except: return ""

def _extraer_de_filas_certificado(filas, candidatos, fechas, bonus=0, origen="tabla"):
    for fila_idx, fila in enumerate(filas):
        if not any(fila): continue
        for columna, celda in enumerate(fila):
            if not celda: continue

            if _contiene_alguna_etiqueta(celda, _ETIQUETAS_NOMBRE_CERTIFICADO):
                _agregar_candidato_especial(candidatos, "nombre", 430 + bonus, _extraer_valor_inline(celda, _ETIQUETAS_NOMBRE_CERTIFICADO), f"{origen}: nombre en misma celda")
                derecha, distancia = _buscar_valor_derecha(fila, columna, "nombre")
                _agregar_candidato_especial(candidatos, "nombre", 420 - distancia + bonus, derecha, f"{origen}: nombre a la derecha")
                debajo, salto = _buscar_valor_debajo(filas, fila_idx, columna, "nombre")
                _agregar_candidato_especial(candidatos, "nombre", 425 - salto + bonus, debajo, f"{origen}: nombre debajo")

            if _contiene_alguna_etiqueta(celda, _ETIQUETAS_CARGO_CERTIFICADO):
                _agregar_candidato_especial(candidatos, "cargo", 420 + bonus, _extraer_valor_inline(celda, _ETIQUETAS_CARGO_CERTIFICADO), f"{origen}: cargo en misma celda")
                derecha, distancia = _buscar_valor_derecha(fila, columna, "cargo")
                _agregar_candidato_especial(candidatos, "cargo", 410 - distancia + bonus, derecha, f"{origen}: cargo a la derecha")
                debajo, salto = _buscar_valor_debajo(filas, fila_idx, columna, "cargo")
                _agregar_candidato_especial(candidatos, "cargo", 415 - salto + bonus, debajo, f"{origen}: cargo debajo")

            if _contiene_alguna_etiqueta(celda, _ETIQUETAS_FECHA_CERTIFICADO) and "NACIMIENTO" not in str(celda).upper():
                fecha_inline = _buscar_fecha_en_texto(celda)
                if fecha_inline: fechas.append((440 + bonus, fecha_inline, f"{origen}: fecha misma celda"))
                fecha_tabla, puntaje = _buscar_fecha_tabla(filas, fila_idx, columna)
                if fecha_tabla: fechas.append((puntaje + bonus, fecha_tabla, f"{origen}: fecha por componentes"))

            if _contiene_alguna_etiqueta(celda, _ETIQUETAS_LUGAR_CERTIFICADO):
                _agregar_candidato_especial(candidatos, "lugar", 435 + bonus, _quitar_fecha_para_lugar(_extraer_valor_inline(celda, _ETIQUETAS_LUGAR_CERTIFICADO)), f"{origen}: lugar misma celda")
                derecha, distancia = _buscar_valor_derecha(fila, columna, "lugar")
                _agregar_candidato_especial(candidatos, "lugar", 425 - distancia + bonus, _quitar_fecha_para_lugar(derecha), f"{origen}: lugar a la derecha")
                debajo, salto = _buscar_valor_debajo(filas, fila_idx, columna, "lugar")
                _agregar_candidato_especial(candidatos, "lugar", 430 - salto + bonus, _quitar_fecha_para_lugar(debajo), f"{origen}: lugar debajo")

            if _contiene_alguna_etiqueta(celda, _ETIQUETAS_IPS_CERTIFICADO):
                _agregar_candidato_especial(candidatos, "lugar", 260 + bonus, _extraer_valor_inline(celda, _ETIQUETAS_IPS_CERTIFICADO), f"{origen}: IPS explícita")

def _filas_desde_lineas(lineas):
    filas = []
    for linea in lineas:
        linea = str(linea).strip()
        if not linea: continue
        columnas = [_limpiar_celda_certificado(col) for col in re.split(r"\s{3,}|\t+|\|", linea)]
        columnas = [c for c in columnas if c]
        if columnas: filas.append(columnas)
    return filas

def extraer_metadatos_pdf_estructurados(pdf_raw_data, texto_completo=""):
    candidatos = {"nombre": [], "cargo": [], "lugar": []}
    fechas = []

    try: documento = pdfplumber.open(io.BytesIO(pdf_raw_data))
    except Exception: documento = None

    if documento is not None:
        with documento as p_file:
            for numero_pagina, page in enumerate(p_file.pages, start=1):
                bonus = max(0, 35 - (numero_pagina - 1) * 8)
                try: tablas = page.extract_tables() or []
                except Exception: tablas = []

                for tabla in tablas:
                    filas = [_normalizar_lista_celdas(fila) for fila in (tabla or [])]
                    _extraer_de_filas_certificado(filas, candidatos, fechas, bonus=bonus, origen=f"tabla página {numero_pagina}")

                lineas_coordenadas = _extraer_lineas_por_coordenadas(page)
                filas_coordenadas = _filas_desde_lineas(lineas_coordenadas)
                _extraer_de_filas_certificado(filas_coordenadas, candidatos, fechas, bonus=bonus - 20, origen=f"coordenadas página {numero_pagina}")

                texto_base = page.extract_text(layout=False) or ""
                if numero_pagina <= 2 or len(re.sub(r"\s+", "", texto_base)) < 150:
                    texto_ocr = _ocr_pagina_si_disponible(page)
                    if texto_ocr:
                        filas_ocr = _filas_desde_lineas(texto_ocr.splitlines())
                        _extraer_de_filas_certificado(filas_ocr, candidatos, fechas, bonus=bonus - 35, origen=f"OCR página {numero_pagina}")
                        fecha_ocr = _buscar_fecha_en_texto(texto_ocr)
                        if fecha_ocr: fechas.append((250 + bonus, fecha_ocr, f"OCR página {numero_pagina}"))

    if texto_completo:
        respaldo = extraer_identidad_cargo_lugar(texto_completo)
        _agregar_candidato_especial(candidatos, "nombre", 170, respaldo.get("nombre", ""), "analizador textual existente")
        _agregar_candidato_especial(candidatos, "cargo", 165, respaldo.get("cargo", ""), "analizador textual existente")
        _agregar_candidato_especial(candidatos, "lugar", 175, respaldo.get("lugar", ""), "analizador textual existente")
        if respaldo.get("fecha"): fechas.append((165, respaldo["fecha"], "analizador textual existente"))

    resultado = {
        "nombre": elegir_mejor_candidato(candidatos["nombre"], "nombre"),
        "cargo": elegir_mejor_candidato(candidatos["cargo"], "cargo"),
        "lugar": elegir_mejor_candidato(candidatos["lugar"], "lugar"),
    }

    fechas_validas = [(puntaje, f, org) for puntaje, f, org in fechas if isinstance(f, datetime.date) and 2000 <= f.year <= 2100]
    if fechas_validas: resultado["fecha"] = max(fechas_validas, key=lambda item: item[0])[1]
    return resultado

def extraer_texto_pdf_robusto(pdf_raw_data):
    lineas_salida = []
    vistos = set()

    def agregar(fragmento, permitir_repetido=False):
        if not fragmento: return
        for linea in str(fragmento).splitlines():
            linea = linea.replace("\t", "    ").strip()
            if not linea: continue
            clave = normalizar_etiqueta(linea)
            if not clave: continue
            if permitir_repetido or clave not in vistos:
                vistos.add(clave)
                lineas_salida.append(linea)

    with pdfplumber.open(io.BytesIO(pdf_raw_data)) as p_file:
        for numero_pagina, page in enumerate(p_file.pages, start=1):
            agregar(page.extract_text(x_tolerance=2, y_tolerance=3, layout=True) or "")
            agregar(page.extract_text(x_tolerance=2, y_tolerance=3, layout=False) or "")
            for linea in _extraer_lineas_por_coordenadas(page): agregar(linea)

            try: tablas = page.extract_tables() or []
            except Exception: tablas = []

            for tabla in tablas:
                for fila in tabla or []:
                    celdas = [re.sub(r"\s+", " ", (celda or "").replace("\n", " ")).strip() for celda in (fila or [])]
                    if any(celdas): agregar(" | ".join(celdas), permitir_repetido=True)

            texto_existente = page.extract_text(layout=False) or ""
            if numero_pagina <= 2 or len(re.sub(r"\s+", "", texto_existente)) < 120:
                texto_ocr = _ocr_pagina_si_disponible(page)
                if texto_ocr: agregar(texto_ocr)

    return "\n".join(lineas_salida)

# --- DETECTOR ESPECÍFICO DE DATO PRIORITARIO PARA CARGO ---
def extraer_cargo_especifico(texto):
    if not texto: return ""
    
    # Expresión regular mejorada para capturar el cargo independientemente del espacio horizontal
    patrones_cargo = [
        r'(?i)\bCargo\b\s*[:=\t|-]*\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9\s/.-]{2,60})',
        r'(?i)\bOcupaci[oó]n\b\s*[:=\t|-]*\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9\s/.-]{2,60})',
        r'(?i)\bPuesto\b\s*[:=\t|-]*\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9\s/.-]{2,60})'
    ]
    
    for p in patrones_cargo:
        m = re.search(p, texto)
        if m:
            candidato = m.group(1).strip()
            # Freno de corte inmediato en palabras reservadas
            candidato = re.split(r'(?i)\b(?:EPS|ARL|AFP|Empresa|Escolaridad|Estado|Tipo|Tipo\s+de\s+Examen|Evaluaci[oó]n|Per[ií]odico|Identificaci[oó]n|Tel[eé]fono|C[eé]dula|Documento)\b', candidato)[0].strip()
            candidato = re.sub(r'\s+', ' ', candidato).strip(" :-,_./|")
            if candidato_cargo_valido(candidato):
                return corregir_ortografia_sst(candidato).title()
    return ""

def extraer_metadatos_formatos_conocidos(texto_completo):
    meta = {}
    
    # Prioridad 1: Extractor de Cargo Robusto
    cargo_directo = extraer_cargo_especifico(texto_completo)
    if cargo_directo:
        meta["cargo"] = cargo_directo

    # 1. FORMATO A: "Fecha y Lugar: 03 jun. 2026 - TUNJA - BOYACA" / "Paciente: MARIA..."
    m_fyl = re.search(r'Fecha\s+y\s+Lugar:\s*(\d{1,2}\s+[a-zA-Z]{3,4}\.?\s+20\d{2})\s*-\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ\s-]+)', texto_completo, re.IGNORECASE)
    if m_fyl:
        f_str, l_str = m_fyl.group(1).strip(), m_fyl.group(2).strip()
        meta["lugar"] = re.split(r'-', l_str)[0].strip().title()
        meta["fecha"] = intentar_parsear_fecha(f_str)
        
    m_pac = re.search(r'Paciente:\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ\s]+?)(?=\s+(?:Identificaci[oó]n|Tel[eé]fono|M[oó]vil|G[eé]nero|Edad|C\.?C|CC)|$)', texto_completo, re.IGNORECASE)
    if m_pac:
        nc = m_pac.group(1).strip()
        if len(nc.split()) >= 2 and candidato_nombre_valido(nc):
            meta["nombre"] = nc.title()

    # 2. FORMATO B: "Apellidos y Nombres" / Grillas
    if "nombre" not in meta:
        m_nom_b = re.search(r'Apellidos\s+y\s+Nombres\s*[:\n|]?\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ\s]+?)(?=\s+(?:G[eé]nero|Edad|Documento|CC|C\.C)|$)', texto_completo, re.IGNORECASE)
        if m_nom_b:
            nc = m_nom_b.group(1).strip()
            if len(nc.split()) >= 2 and candidato_nombre_valido(nc):
                meta["nombre"] = nc.title()

    if "cargo" not in meta:
        lines = texto_completo.splitlines()
        for idx, line in enumerate(lines):
            if line.strip().upper() == "CARGO" and idx + 1 < len(lines):
                next_l = lines[idx+1].strip()
                col1 = re.split(r'\s{2,}|\|', next_l)[0].strip()
                if candidato_cargo_valido(col1):
                    meta["cargo"] = corregir_ortografia_sst(col1).title()
                    break

    if "fecha" not in meta or "lugar" not in meta:
        m_grid_b = re.search(r'\b(\d{1,2})\s*[\s|/-]+\s*(\d{1,2})\s*[\s|/-]+\s*(20\d{2})\s*\|\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ\s]+?)(?=\(|\n|$)', texto_completo)
        if m_grid_b:
            d, m, y = int(m_grid_b.group(1)), int(m_grid_b.group(2)), int(m_grid_b.group(3))
            lug = m_grid_b.group(4).strip()
            if 1 <= d <= 31 and 1 <= m <= 12: meta["fecha"] = datetime.date(y, m, d)
            if candidato_lugar_valido(lug): meta["lugar"] = lug.title()

    return meta

def deduplicar_textos(items):
    resultado = []
    for item in items or []:
        limpio = re.sub(r"\s+", " ", str(item or "")).strip(" •\t\r\n-_:;.,")
        if not limpio:
            continue
        norm = normalizar_etiqueta(limpio)
        reemplazado = False
        for idx, existente in enumerate(resultado):
            norm_existente = normalizar_etiqueta(existente)
            if norm == norm_existente:
                reemplazado = True
                break
            tokens_a, tokens_b = set(norm.split()), set(norm_existente.split())
            similitud = len(tokens_a & tokens_b) / max(1, len(tokens_a | tokens_b))
            if similitud >= 0.78 or norm in norm_existente or norm_existente in norm:
                if len(limpio) > len(existente):
                    resultado[idx] = limpio
                reemplazado = True
                break
        if not reemplazado:
            resultado.append(limpio)
    return resultado

def extraer_examenes_globales(texto):
    """Recupera exámenes aunque la tabla PDF pierda columnas o saltos de línea."""
    encontrados = []
    dentro_seccion = False
    encabezados_examen = ("EXAMENES REALIZADOS", "EXAMENES PRACTICADOS", "PRUEBAS REALIZADAS", "EXAMENES EFECTUADOS")
    encabezados_salida = ("RECOMENDACIONES", "OBSERVACIONES", "REMISIONES", "VIGILANCIA", "CONSENTIMIENTO")
    verbos_futuros = ("REALIZAR", "SOLICITAR", "REMITIR", "PROGRAMAR", "ORDENAR", "ASISTIR A", "CONTROL POR")
    for linea in texto.splitlines():
        linea_norm = normalizar_etiqueta(linea)
        if any(encabezado in linea_norm for encabezado in encabezados_examen):
            dentro_seccion = True
            continue
        if dentro_seccion and any(encabezado in linea_norm for encabezado in encabezados_salida):
            dentro_seccion = False
        for clave in sorted(EXAMS_MAP, key=len, reverse=True):
            clave_norm = normalizar_etiqueta(clave)
            coincidencia = re.search(rf"(?<![A-Z0-9]){re.escape(clave_norm)}(?![A-Z0-9])", linea_norm)
            if not coincidencia:
                continue
            prefijo = linea_norm[:coincidencia.start()]
            sugerido_para_futuro = any(verbo in prefijo for verbo in verbos_futuros)
            aparece_como_item = coincidencia.start() <= 18 or "REALIZADO" in linea_norm or "PRACTICADO" in linea_norm
            if (dentro_seccion or aparece_como_item) and not sugerido_para_futuro:
                valor = EXAMS_MAP[clave]
                if valor not in encontrados:
                    encontrados.append(valor)
    return encontrados

def extraer_recomendaciones_genericas(texto):
    """Respaldo para recomendaciones envueltas en varias líneas o columnas."""
    encabezados_inicio = (
        "RECOMENDACIONES MEDICAS", "RECOMENDACIONES OCUPACIONALES",
        "RECOMENDACIONES GENERALES", "HABITOS Y ESTILO DE VIDA SALUDABLES"
    )
    encabezados_fin = (
        "OTRAS OBSERVACIONES", "OBSERVACIONES:", "INFORMACION DE REMISIONES",
        "INFORMACIÓN DE REMISIONES", "REMISIONES:", "CONSENTIMIENTO",
        "AUTORIZO", "TRATAMIENTO DE DATOS", "FIRMA DEL TRABAJADOR", "ATENTAMENTE"
    )
    dentro = False
    fragmentos = []
    actual = ""
    for linea in texto.splitlines():
        linea = limpiar_linea_ruido_lateral(linea)
        norm = normalizar_etiqueta(linea)
        if not dentro and any(h in norm for h in encabezados_inicio):
            dentro = True
            continue
        if dentro and any(h in norm for h in encabezados_fin):
            break
        if not dentro or not linea.strip():
            continue
        columnas = [c.strip() for c in re.split(r"\s{3,}|\|", linea) if c.strip()]
        for columna in columnas:
            candidato = recortar_contenido_legal(columna)
            candidato = re.sub(r"^\s*(?:[•*-]|\d+[.)-])\s*", "", candidato).strip()
            if es_vacio_o_estado(candidato) or es_contenido_legal_recomendacion(candidato):
                continue
            comienza_nueva = bool(re.match(
                r"^(?:REALIZAR|ASISTIR|UTILIZAR|USAR|MANTENER|EVITAR|CONTINUAR|CONTROLAR|SOLICITAR|SEGUIR)\b",
                normalizar_etiqueta(candidato)
            ))
            if actual and not re.search(r"[.!?;:]$", actual) and not comienza_nueva:
                actual = f"{actual} {candidato}"
            else:
                if actual:
                    fragmentos.append(a_caso_oracion(actual))
                actual = candidato
    if actual:
        fragmentos.append(a_caso_oracion(actual))
    return deduplicar_textos(fragmentos)

def obtener_gemini_api_key():
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return str(st.secrets["GEMINI_API_KEY"]).strip()
    except Exception:
        pass
    return os.getenv("GEMINI_API_KEY", "").strip() or obtener_config("gemini_api_key").strip()

def extraer_identificacion_correo(texto):
    """Recupera cédula y correo sin confundirlos con NIT, teléfonos o consecutivos."""
    resultado = {"identificacion": "", "correo": ""}
    texto = str(texto or "")
    correo = re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", texto, re.IGNORECASE)
    if correo:
        resultado["correo"] = correo.group(0).lower()

    patrones_documento = [
        r"(?:N[ÚU]MERO\s+DE\s+)?(?:DOCUMENTO|IDENTIFICACI[ÓO]N|C[ÉE]DULA)(?:\s+DE\s+CIUDADAN[IÍ]A)?\s*[:#-]?\s*(?:C\.?\s*C\.?\s*)?([0-9][0-9.\s-]{5,17})",
        r"\bC\.?\s*C\.?\s*(?:N[ÚU]MERO|NO\.?|N[°º])?\s*[:#-]?\s*([0-9][0-9.\s-]{5,17})"
    ]
    for patron in patrones_documento:
        coincidencia = re.search(patron, texto, re.IGNORECASE)
        if not coincidencia:
            continue
        numero = re.sub(r"\D", "", coincidencia.group(1))
        if 6 <= len(numero) <= 12:
            resultado["identificacion"] = numero
            break
    return resultado

def _texto_respaldado_por_fuente(valor, texto_fuente, umbral=0.42):
    valor_norm = normalizar_etiqueta(valor)
    fuente_norm = normalizar_etiqueta(texto_fuente)
    if not valor_norm:
        return False
    if valor_norm in fuente_norm:
        return True
    palabras = [p for p in valor_norm.split() if len(p) >= 4]
    if not palabras:
        return valor_norm in fuente_norm
    coincidencias = sum(1 for p in palabras if p in fuente_norm)
    return coincidencias / len(palabras) >= umbral

def _extraer_json_interaccion_gemini(respuesta):
    respuesta.raise_for_status()
    cuerpo = respuesta.json()
    fragmentos = []
    for paso in cuerpo.get("steps", []):
        if paso.get("type") == "model_output":
            fragmentos.extend(
                contenido.get("text", "")
                for contenido in paso.get("content", [])
                if contenido.get("type") == "text" or contenido.get("text")
            )
    if not fragmentos:
        fragmentos.extend(
            salida.get("text", "") for salida in cuerpo.get("outputs", [])
            if salida.get("type") == "text" or salida.get("text")
        )
    texto_json = "".join(fragmentos).strip()
    texto_json = re.sub(r"^```(?:json)?\s*|\s*```$", "", texto_json, flags=re.IGNORECASE)
    if not texto_json:
        raise ValueError("Gemini respondió, pero no devolvió el JSON de extracción.")
    return json.loads(texto_json)

def _crear_payload_gemini_pdf(pdf_bytes, prompt, esquema, modelo):
    return {
        "model": modelo,
        "input": [
            {
                "type": "document",
                "data": base64.b64encode(pdf_bytes).decode("ascii"),
                "mime_type": "application/pdf"
            },
            {"type": "text", "text": prompt}
        ],
        "response_format": {
            "type": "text", "mime_type": "application/json", "schema": esquema
        },
        "generation_config": {"temperature": 0}
    }

def validar_documento_con_gemini(pdf_bytes, texto, datos_locales, api_key, modelo=GEMINI_MODEL_DEFAULT):
    """Lee el PDF visual completo con Interactions API y devuelve una transcripción estructurada."""
    if not api_key:
        raise ValueError("Configura una clave de Gemini antes de usar la validación por IA.")
    if not pdf_bytes:
        raise ValueError("No se encontraron los bytes del PDF para la validación visual.")
    if len(pdf_bytes) > 50 * 1024 * 1024:
        raise ValueError("El PDF supera el límite de 50 MB admitido para lectura directa.")
    esquema = {
        "type": "object",
        "properties": {
            "nombre": {"type": "string"}, "cargo": {"type": "string"},
            "identificacion": {"type": "string"}, "correo": {"type": "string"},
            "tipo_examen": {"type": "string"}, "lugar": {"type": "string"},
            "fecha": {"type": "string", "description": "AAAA-MM-DD o vacío"},
            "examenes_realizados": {"type": "array", "items": {"type": "string"}},
            "recomendaciones_medicas": {"type": "array", "items": {"type": "string"}},
            "recomendaciones_por_examen": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "examen": {"type": "string"},
                        "recomendaciones": {"type": "array", "items": {"type": "string"}}
                    },
                    "required": ["examen", "recomendaciones"]
                }
            },
            "vigilancia_programa": {"type": "array", "items": {"type": "string"}},
            "observaciones": {"type": "string"}, "remisiones": {"type": "string"}
        },
        "required": ["nombre", "cargo", "identificacion", "correo", "tipo_examen", "lugar", "fecha", "examenes_realizados",
                     "recomendaciones_medicas", "recomendaciones_por_examen", "vigilancia_programa", "observaciones", "remisiones"]
    }
    prompt = f"""Eres un extractor documental para Seguridad y Salud en el Trabajo.
Lee visualmente TODAS las páginas del PDF adjunto, incluidas tablas, columnas, celdas y textos escaneados.
Devuelve únicamente los datos presentes en el documento.
Reglas obligatorias:
- No diagnostiques, no recomiendes y no inventes información.
- Transcribe cada recomendación COMPLETA desde su inicio hasta su punto final, aunque continúe en otra línea o celda.
- Une correctamente los saltos de línea que pertenecen a una misma recomendación.
- No resumas, no parafrasees y no cortes frases.
- Convierte los bloques escritos totalmente en mayúsculas a redacción normal en minúsculas, usando mayúscula solo al iniciar cada oración y en nombres propios o siglas.
- Corrige ortografía, tildes, espacios y puntuación evidentes sin añadir, eliminar ni cambiar el sentido del texto.
- No dupliques exámenes ni recomendaciones; devuelve cada elemento una sola vez.
- Relaciona cada recomendación con el examen que la origina en recomendaciones_por_examen.
- Incluye en recomendaciones_por_examen todos los exámenes realizados. Si un examen no tiene una recomendación explícita, usa una lista vacía; no inventes texto.
- Usa recomendaciones_medicas como lista plana de respaldo, conservando el prefijo «Examen:» cuando la relación sea explícita.
- Extrae el número de identificación del trabajador y su correo electrónico cuando estén explícitos en el PDF.
- Enumera todos los exámenes explícitamente realizados; no incluyas exámenes apenas sugeridos para el futuro.
- Omite consentimientos, habeas data, firmas y texto legal.
- Si un dato no aparece, usa cadena vacía o lista vacía.
- Corrige únicamente errores OCR evidentes sin cambiar el significado clínico.
- Revisa el PDF completo una segunda vez antes de responder y confirma que ningún examen o recomendación quedó omitido.

Extracción local de referencia (puede estar incompleta):
{json.dumps(datos_locales, ensure_ascii=False, default=str)}

TEXTO EXTRAÍDO LOCALMENTE COMO APOYO:
{texto[:30000]}"""

    modelos_candidatos = []
    for candidato in [modelo, GEMINI_MODEL_DEFAULT, "gemini-3.1-flash-lite"]:
        candidato = str(candidato or "").strip().removeprefix("models/")
        if candidato and candidato not in modelos_candidatos:
            modelos_candidatos.append(candidato)

    ultimo_error = None
    for modelo_candidato in modelos_candidatos:
        payload = _crear_payload_gemini_pdf(pdf_bytes, prompt, esquema, modelo_candidato)
        try:
            respuesta = requests.post(
                GEMINI_INTERACTIONS_URL,
                headers={
                    "x-goog-api-key": api_key,
                    "Content-Type": "application/json",
                    "Api-Revision": "2026-05-20"
                },
                json=payload,
                timeout=150
            )
            if respuesta.status_code == 404:
                ultimo_error = f"El modelo {modelo_candidato} no está disponible para esta clave."
                continue
            datos = _extraer_json_interaccion_gemini(respuesta)
            mapa_ia = agrupar_recomendaciones_por_examen(
                datos.get("examenes_realizados", []),
                datos.get("recomendaciones_medicas", []),
                datos.get("recomendaciones_por_examen", [])
            )
            recomendaciones, pendientes = separar_recomendaciones_atomicas(
                aplanar_recomendaciones_por_examen(mapa_ia)
            )
            datos["recomendaciones_por_examen"] = mapa_ia
            locales_atomicas, _ = separar_recomendaciones_atomicas(
                datos_locales.get("recomendaciones_lista", [])
            )
            datos["recomendaciones_medicas"] = recomendaciones
            necesita_revision = bool(pendientes) or (
                len(locales_atomicas) >= 3 and len(recomendaciones) < max(1, len(locales_atomicas) // 2)
            )
            datos["_segunda_revision_ia"] = False

            if necesita_revision:
                prompt_revision = f"""Realiza una segunda auditoría visual completa del PDF adjunto.
La primera lectura produjo el siguiente JSON:
{json.dumps(datos, ensure_ascii=False, default=str)}

Se detectaron fragmentos posiblemente cortados o información insuficiente:
{json.dumps(pendientes, ensure_ascii=False)}

Devuelve de nuevo TODO el JSON del certificado, no solo las correcciones.
Reglas obligatorias:
- Lee todas las páginas, tablas y continuaciones de renglón.
- Cada elemento de recomendaciones_medicas debe contener UNA sola recomendación completa.
- Si una recomendación pertenece a un examen, usa el formato «Nombre del examen: recomendación completa».
- Completa recomendaciones_por_examen con un registro para cada examen realizado y conserva listas vacías cuando el PDF no asocie una recomendación explícita.
- No uses numeraciones, barras dobles ni varias recomendaciones dentro del mismo elemento.
- No completes por suposición: reconstruye únicamente con texto comprobable en el PDF.
- No repitas elementos y no cambies el sentido clínico.
- Si una frase realmente no puede leerse completa, no la inventes y déjala fuera.
"""
                try:
                    respuesta_revision = requests.post(
                        GEMINI_INTERACTIONS_URL,
                        headers={
                            "x-goog-api-key": api_key,
                            "Content-Type": "application/json",
                            "Api-Revision": "2026-05-20"
                        },
                        json=_crear_payload_gemini_pdf(
                            pdf_bytes, prompt_revision, esquema, modelo_candidato
                        ),
                        timeout=150
                    )
                    if respuesta_revision.ok:
                        datos_revision = _extraer_json_interaccion_gemini(respuesta_revision)
                        mapa_revision = agrupar_recomendaciones_por_examen(
                            datos_revision.get("examenes_realizados", []),
                            datos_revision.get("recomendaciones_medicas", []),
                            datos_revision.get("recomendaciones_por_examen", [])
                        )
                        recomendaciones_revision, pendientes_revision = separar_recomendaciones_atomicas(
                            aplanar_recomendaciones_por_examen(mapa_revision)
                        )
                        calidad_primera = (len(pendientes), -len(recomendaciones))
                        calidad_revision = (len(pendientes_revision), -len(recomendaciones_revision))
                        if calidad_revision <= calidad_primera:
                            datos = datos_revision
                            recomendaciones = recomendaciones_revision
                            pendientes = pendientes_revision
                            datos["recomendaciones_por_examen"] = mapa_revision
                        datos["_segunda_revision_ia"] = True
                except (requests.RequestException, ValueError, json.JSONDecodeError):
                    datos["_segunda_revision_ia"] = False

            datos["recomendaciones_medicas"] = recomendaciones
            datos["recomendaciones_por_examen"] = agrupar_recomendaciones_por_examen(
                datos.get("examenes_realizados", []), recomendaciones,
                datos.get("recomendaciones_por_examen", [])
            )
            datos["_fragmentos_pendientes"] = pendientes
            datos["_modelo_usado"] = modelo_candidato
            return datos
        except requests.RequestException as exc:
            ultimo_error = f"{type(exc).__name__}: {exc}"
            if getattr(exc.response, "status_code", None) not in {404, 429, 503}:
                break
        except (ValueError, json.JSONDecodeError) as exc:
            ultimo_error = str(exc)
            break
    raise RuntimeError(ultimo_error or "Gemini no devolvió una extracción utilizable.")

def fusionar_validacion_ia(datos_locales, datos_ia, texto_fuente):
    """Prioriza la lectura visual completa del PDF; conserva el extractor local como respaldo."""
    resultado = dict(datos_locales)
    mapeo = {
        "nombre": "nombre", "cargo": "cargo", "tipo_examen": "tipo_examen",
        "lugar": "lugar", "identificacion": "identificacion", "correo": "correo",
        "observaciones": "observaciones", "remisiones": "remisiones"
    }
    for destino, origen in mapeo.items():
        candidato = str(datos_ia.get(origen, "") or "").strip()
        if candidato:
            resultado[destino] = a_caso_oracion(candidato) if destino in {"observaciones", "remisiones"} else candidato

    examenes_ia = deduplicar_textos(datos_ia.get("examenes_realizados", []))
    recomendaciones_ia = filtrar_recomendaciones_clinicas(
        deduplicar_textos(datos_ia.get("recomendaciones_medicas", []))
    )
    if examenes_ia:
        resultado["examenes_lista"] = examenes_ia
    if recomendaciones_ia:
        resultado["recomendaciones_lista"] = recomendaciones_ia
    mapa_ia = agrupar_recomendaciones_por_examen(
        examenes_ia or resultado.get("examenes_lista", []),
        recomendaciones_ia,
        datos_ia.get("recomendaciones_por_examen", [])
    )
    if any(mapa_ia.values()):
        resultado["recomendaciones_por_examen"] = mapa_ia
    resultado["recomendaciones_pendientes_revision"] = deduplicar_textos(
        datos_ia.get("_fragmentos_pendientes", [])
    )
    programas = deduplicar_textos(datos_ia.get("vigilancia_programa", []))
    if programas:
        resultado["vigilancia_programa"] = ", ".join(programas)

    fecha_ia = str(datos_ia.get("fecha", "") or "").strip()
    if fecha_ia and re.fullmatch(r"20\d{2}-\d{2}-\d{2}", fecha_ia):
        try:
            resultado["fecha"] = datetime.datetime.strptime(fecha_ia, "%Y-%m-%d").date()
        except ValueError:
            pass
    resultado["validado_ia"] = True
    resultado["modelo_ia"] = datos_ia.get("_modelo_usado", "")
    resultado["segunda_revision_ia"] = bool(datos_ia.get("_segunda_revision_ia"))
    resultado["modo_validacion"] = (
        "IA automática + segunda revisión de calidad"
        if resultado["segunda_revision_ia"] else "IA automática + respaldo local"
    )
    return normalizar_datos_documento(resultado)

# --- ANALIZADOR INTELIGENTE DE DOCUMENTOS ---
def analizar_pdf_inteligente(texto, metadatos_pdf=None):
    datos = {
        "nombre": "", "cargo": "", "tipo_examen": "PERIODICO",
        "identificacion": "", "correo": "",
        "examenes_lista": [], "recomendaciones_lista": [], "vigilancia_lista": [],
        "observaciones": "", "remisiones": "No", "consecutivo": "",
        "vigilancia_programa": "NINGUNO", "lugar": "Tunja", "fecha": datetime.date.today()
    }
    if not texto: return datos

    contacto = extraer_identificacion_correo(texto)
    datos.update({clave: valor for clave, valor in contacto.items() if valor})

    # 1. Extracción de Alta Prioridad
    meta_conocida = extraer_metadatos_formatos_conocidos(texto)
    for k, v in meta_conocida.items():
        if v: datos[k] = v

    # 2. Respaldo por Extractor Matricial
    identificacion = extraer_identidad_cargo_lugar(texto)
    for k in ["nombre", "cargo", "lugar", "fecha"]:
        if not datos[k] and identificacion.get(k):
            datos[k] = identificacion[k]

    metadatos_pdf = metadatos_pdf or {}
    for k in ["nombre", "cargo", "lugar", "fecha"]:
        if not datos[k] and metadatos_pdf.get(k):
            datos[k] = metadatos_pdf[k]

    for palabra in ["INGRESO", "PERIÓDICO", "PERIODICO", "EGRESO", "RETIRO", "CAMBIO DE CARGO", "POST-INCAPACIDAD", "POST INCAPACIDAD", "CONTROL PERIÓDICO"]:
        if palabra in texto.upper():
            datos["tipo_examen"] = "PERIODICO" if "PERIOD" in palabra or "CONTROL" in palabra else palabra
            break

    lineas_raw = texto.split("\n")
    examenes_detectados = []
    recoms_raw_dict = {}
    current_exam = None
    in_exams_section = True
    formato_grilla_detectado = False
    recoms_grilla_acumuladas = []

    for idx_l, linea in enumerate(lineas_raw):
        linea_limpia = limpiar_linea_ruido_lateral(linea)
        linea_upper = linea_limpia.upper().strip()
        
        if "EL CONCEPTO DE APTITUD SE DEFINIÓ A PARTIR DE LOS SIGUIENTES EXÁMENES PRACTICADOS" in linea_upper:
            for offset in range(1, 4):
                if idx_l + offset < len(lineas_raw):
                    l_sig = lineas_raw[idx_l + offset].upper()
                    for k_ex, v_ex in EXAMS_MAP.items():
                        if k_ex in l_sig and v_ex not in examenes_detectados: examenes_detectados.append(v_ex)
            continue
            
        if any(h in linea_upper for h in ["RECOMENDACIONES MÉDICAS", "RECOMENDACIONES MEDICAS", "RECOMENDACIONES OCUPACIONALES", "HÁBITOS Y ESTILO DE VIDA SALUDABLES", "HABITOS Y ESTILO DE VIDA SALUDABLES"]):
            formato_grilla_detectado = True
            current_exam = None
            continue
            
        if formato_grilla_detectado:
            if any(stop in linea_upper for stop in ["OTRAS OBSERVACIONES", "INFORMACION DE REMISIONES", "INFORMACIÓN DE REMISIONES", "REMISIONES", "ATENTAMENTE", "CONSENTIMIENTO", "AUTORIZO", "TRATAMIENTO DE DATOS", "HABEAS DATA", "FIRMA DEL TRABAJADOR"]) or es_encabezado_legal(linea_limpia) or es_contenido_legal_recomendacion(linea_limpia):
                formato_grilla_detectado = False
                current_exam = None
                continue
            else:
                columnas = [col.strip(" |/-,_.") for col in re.split(r'\s{2,}|\|', linea_limpia) if col.strip()]
                for col in columnas:
                    col_clinica = recortar_contenido_legal(col)
                    if col_clinica and not es_vacio_o_estado(col_clinica) and not es_contenido_legal_recomendacion(col_clinica):
                        rec_fmt = a_caso_oracion(col_clinica)
                        if rec_fmt and rec_fmt not in recoms_grilla_acumuladas: recoms_grilla_acumuladas.append(rec_fmt)
                continue

        if any(stop in linea_upper for stop in ["OTRAS OBSERVACIONES", "OBSERVACIONES", "OBSERVACION", "INFORMACION DE REMISIONES", "INFORMACIÓN DE REMISIONES", "REMISIONES", "SISTEMA DE VIGILANCIA", "CONSENTIMIENTO", "AUTORIZO", "ATENTAMENTE"]) or es_encabezado_legal(linea_limpia) or es_contenido_legal_recomendacion(linea_limpia):
            in_exams_section = False
            if current_exam:
                contenido_actual = recoms_raw_dict.get(current_exam, "")
                recoms_raw_dict[current_exam] = recortar_contenido_legal(contenido_actual)
                current_exam = None
            continue

        matched_key = None
        for key in sorted(EXAMS_MAP.keys(), key=len, reverse=True):
            if key in linea_upper and linea_upper.find(key) < 15:
                matched_key = key
                break
        
        if matched_key:
            posicion_examen = linea_upper.find(matched_key)
            prefijo_examen = linea_upper[:posicion_examen]
            if not in_exams_section and any(verbo in prefijo_examen for verbo in ["REALIZAR", "SOLICITAR", "REMITIR", "PROGRAMAR", "ORDENAR", "CONTROL POR"]):
                continue
            in_exams_section = True
            current_exam = EXAMS_MAP[matched_key]
            if current_exam not in examenes_detectados: examenes_detectados.append(current_exam)
            idx = posicion_examen + len(matched_key)
            recoms_raw_dict[current_exam] = linea_limpia[idx:].strip(" :-,_/")
        else:
            if in_exams_section and current_exam and linea_limpia.strip():
                if not es_encabezado_legal(linea_limpia) and not es_contenido_legal_recomendacion(linea_limpia):
                    recoms_raw_dict[current_exam] = recoms_raw_dict.get(current_exam, "") + " " + linea_limpia.strip()

    recoms_por_examen = []
    pve_detectados = set()

    if recoms_grilla_acumuladas:
        for rec in recoms_grilla_acumuladas:
            recoms_por_examen.append(rec)
            rec_up = rec.upper()
            if any(re.search(patron, rec_up) for patron in [r'\bAUDITIV', r'\bRUIDO', r'\bOIDO', r'\bOÍDO', r'\bAUDIO']): pve_detectados.add("Conservación Auditiva")
            if any(re.search(patron, rec_up) for patron in [r'\bPOSTURAL', r'\bLUMBAR', r'\bOSTEOMUSCULAR', r'\bERGONOMIC', r'\bESPALDA', r'\bCARGA']): pve_detectados.add("Prevención Osteomuscular (DME)")
            if any(re.search(patron, rec_up) for patron in [r'\bVISUAL', r'\bGAFAS', r'\bVISION', r'\bVISIÓN', r'\bLENTE', r'\bOPTOMETR', r'\bRX\b']): pve_detectados.add("Conservación Visual")
            if any(re.search(patron, rec_up) for patron in [r'\bRESPIRATORI', r'\bESPIROMETR', r'\bPOLVO', r'\bHUMO']): pve_detectados.add("Conservación Respiratoria")
    else:
        for exam in examenes_detectados:
            rec_part = recoms_raw_dict.get(exam, "").strip()
            rec_part = recortar_contenido_legal(rec_part)
            rec_part = re.sub(r'\s+', ' ', rec_part)
            rec_part = limpiar_ruido_columnas_final(rec_part)
            
            if not es_vacio_o_estado(rec_part):
                parts = re.split(r'//|;|\b\d+\.|\b\d+\-', rec_part)
                valid_parts = []
                for p in parts:
                    p_clean = p.strip(" .-_/()[]")
                    p_clean = recortar_contenido_legal(p_clean)
                    if p_clean and not es_vacio_o_estado(p_clean) and not es_contenido_legal_recomendacion(p_clean):
                        valid_parts.append(a_caso_oracion(p_clean))
                        p_upper = p_clean.upper()
                        if any(re.search(patron, p_upper) for patron in [r'\bAUDITIV', r'\bRUIDO', r'\bOIDO', r'\bOÍDO', r'\bAUDIO']): pve_detectados.add("Conservación Auditiva")
                        if any(re.search(patron, p_upper) for patron in [r'\bPOSTURAL', r'\bLUMBAR', r'\bOSTEOMUSCULAR', r'\bERGONOMIC', r'\bESPALDA', r'\bCARGA']): pve_detectados.add("Prevención Osteomuscular (DME)")
                        if any(re.search(patron, p_upper) for patron in [r'\bVISUAL', r'\bGAFAS', r'\bVISION', r'\bVISIÓN', r'\bLENTE', r'\bOPTOMETR', r'\bRX\b']): pve_detectados.add("Conservación Visual")
                        if any(re.search(patron, p_upper) for patron in [r'\bRESPIRATORI', r'\bESPIROMETR', r'\bPOLVO', r'\bHUMO']): pve_detectados.add("Conservación Respiratoria")
                if valid_parts: recoms_por_examen.append(f"{exam}: {' - '.join(valid_parts)}")

    datos["examenes_lista"] = deduplicar_textos(examenes_detectados + extraer_examenes_globales(texto))
    recomendaciones_respaldo = extraer_recomendaciones_genericas(texto)
    datos["recomendaciones_lista"] = filtrar_recomendaciones_clinicas(
        deduplicar_textos(recoms_por_examen + recomendaciones_respaldo)
    )
    datos["vigilancia_lista"] = list(pve_detectados)

    # RECOLECCIÓN VECINAL ESTRICTA DE PVE (MAX 3 LÍNEAS DEBAJO DE LA CABECERA)
    programas_encontrados = []
    for idx, line in enumerate(lineas_raw):
        l_up = line.upper()
        if "INGRESAR AL PROGRAMA DE VIGILANCIA" in l_up or "PROGRAMA DE VIGILANCIA EPIDEMIOL" in l_up:
            for offset in [0, 1, 2, 3]:
                if idx + offset < len(lineas_raw):
                    text_target = lineas_raw[idx + offset].upper()
                    if offset > 0 and any(stop in text_target for stop in ["REMISIONES:", "OBSERVACIONES:", "ATENTAMENTE", "CONSENTIMIENTO", "AUTORIZO"]): break
                    for kw, prog_name in SVE_CLINICAL_KEYWORDS.items():
                        if kw in text_target and prog_name not in programas_encontrados: programas_encontrados.append(prog_name)
            break

    for pve_bandera in datos["vigilancia_lista"]:
        if pve_bandera not in programas_encontrados: programas_encontrados.append(pve_bandera)
        
    datos["vigilancia_programa"] = ", ".join(programas_encontrados) if programas_encontrados else "NINGUNO"

    def extraer_seccion_limpia(texto_completo, palabras_inicio, palabras_fin):
        seccion = []
        dentro = False
        for l in texto_completo.split('\n'):
            l_limpia = limpiar_linea_ruido_lateral(l)
            l_upper = l_limpia.upper().strip()
            if not dentro:
                if any(h in l_upper for h in palabras_inicio):
                    dentro = True
                    for h in palabras_inicio:
                        if h in l_upper:
                            resto = l_limpia[l_upper.find(h) + len(h):].strip(" :-,_")
                            if resto: seccion.append(resto)
                            break
            else:
                if any(h in l_upper for h in palabras_fin): break
                seccion.append(l_limpia)
        return "\n".join([s for s in seccion if s]).strip()

    obs_fmt_nuevo = ""
    m_obs_nuevo = re.search(r'OTRAS OBSERVACIONES Y RECOMENDACIONES\s*\n\s*([^\n]+)', texto, re.IGNORECASE)
    if m_obs_nuevo: obs_fmt_nuevo = m_obs_nuevo.group(1).strip()
        
    if obs_fmt_nuevo and not es_vacio_o_estado(obs_fmt_nuevo): datos["observaciones"] = a_caso_oracion(obs_fmt_nuevo)
    else: datos["observaciones"] = a_caso_oracion(extraer_seccion_limpia(texto, ["OBSERVACIONES:"], ["RECOMENDACIONES", "REMISIONES", "INGRESAR AL PROGRAMA", "PROGRAMA DE VIGILANCIA"]))
    
    rem_raw = extraer_seccion_limpia(texto, ["INFORMACION DE REMISIONES", "INFORMACIÓN DE REMISIONES"], ["CONSENTIMIENTO", "AUTORIZO"])
    datos["remisiones"] = "No" if es_vacio_o_negativo(rem_raw) else a_caso_oracion(rem_raw)
    datos["modo_validacion"] = "Respaldo local"
    return normalizar_datos_documento(datos)

# --- MOTOR DE RENDERIZADO WORD ---
def aplicar_negrita_dinamica_cuerpo(paragraph, tipo_examen):
    texto_parrafo = paragraph.text
    if "Según los lineamientos del programa de medicina preventiva" not in texto_parrafo:
        return
        
    paragraph.text = "" 
    p1 = "Según los lineamientos del programa de medicina preventiva y del trabajo de JER S.A; se hace entrega de las recomendaciones establecidas por el Proveedor de servicios de Exámenes Médico Ocupacionales ("
    paragraph.add_run(p1)
    
    opciones = [
        ("Ingreso", "INGRESO" in tipo_examen.upper()),
        ("Periódico", "PERIODIC" in tipo_examen.upper() or "PERIÓDIC" in tipo_examen.upper()),
        ("egreso", "EGRESO" in tipo_examen.upper() or "RETIRO" in tipo_examen.upper()),
        ("cambio de cargo", "CAMBIO" in tipo_examen.upper()),
        ("post incapacidad", "POST" in tipo_examen.upper() or "INCAPACIDAD" in tipo_examen.upper())
    ]
    
    for i, (texto_opcion, condicion) in enumerate(opciones):
        run = paragraph.add_run(texto_opcion)
        run.bold = condicion 
        if i < len(opciones) - 1:
            if i == len(opciones) - 2: paragraph.add_run(" y ")
            else: paragraph.add_run(", ")
                
    paragraph.add_run(")")

def replace_placeholder_in_paragraph_runs(paragraph, placeholder, value):
    if placeholder not in paragraph.text: return False
    replaced = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            replaced = True
            
    if not replaced:
        font_name, font_size, bold, italic, color = "Arial", Pt(11), False, False, None
        if paragraph.runs:
            for r in paragraph.runs:
                if r.text.strip():
                    font_name = r.font.name or font_name; font_size = r.font.size or font_size
                    bold = r.bold if r.bold is not None else bold; italic = r.italic if r.italic is not None else italic
                    color = r.font.color.rgb if r.font.color else color
                    break
        full_text = paragraph.text.replace(placeholder, value)
        paragraph.text = ""
        new_run = paragraph.add_run(full_text)
        new_run.font.name = font_name; new_run.font.size = font_size; new_run.bold = bold; new_run.italic = italic
        if color: new_run.font.color.rgb = color
    return True

def insert_bullets_in_placeholder(parent_container, paragraph, items_list):
    if not paragraph.runs: font_name, font_size, bold, color = "Arial", Pt(11), False, None
    else:
        first_run = paragraph.runs[0]
        font_name, font_size, bold, color = first_run.font.name or "Arial", first_run.font.size or Pt(11), first_run.bold, first_run.font.color.rgb if first_run.font.color else None
    paragraph.text = ""; paragraph.paragraph_format.space_after = Pt(2); paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    if not items_list:
        run = paragraph.add_run("Ninguno."); run.font.name = font_name; run.font.size = font_size; run.bold = bold
        if color: run.font.color.rgb = color
        return
    run = paragraph.add_run("• " + items_list[0]); run.font.name = font_name; run.font.size = font_size; run.bold = bold
    if color: run.font.color.rgb = color
    current_p = paragraph
    for item in items_list[1:]:
        new_p_element = OxmlElement('w:p'); current_p._p.addnext(new_p_element)
        new_para = Paragraph(new_p_element, parent_container)
        new_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
        new_para.paragraph_format.space_after = Pt(2); new_para.paragraph_format.space_before = Pt(0)
        new_para.paragraph_format.left_indent = Inches(0.28)
        new_para.paragraph_format.first_line_indent = Inches(-0.18)
        run_new = new_para.add_run("• " + item); run_new.font.name = font_name; run_new.font.size = font_size; run_new.bold = bold
        if color: run_new.font.color.rgb = color
        current_p = new_para

def insert_recommendations_in_placeholder(parent_container, paragraph, datos):
    mapa = agrupar_recomendaciones_por_examen(
        datos.get("examenes_lista", []),
        datos.get("recomendaciones_lista", []),
        datos.get("recomendaciones_por_examen", {})
    )
    if paragraph.runs: font_name, font_size, color = paragraph.runs[0].font.name or "Arial", paragraph.runs[0].font.size or Pt(11), paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color else None
    else: font_name, font_size, color = "Arial", Pt(11), None
    paragraph.text = ""; paragraph.paragraph_format.space_after = Pt(4); paragraph.paragraph_format.space_before = Pt(0)
    run_lbl = paragraph.add_run("Recomendaciones: "); run_lbl.bold = True; run_lbl.font.name = font_name; run_lbl.font.size = font_size
    if color: run_lbl.font.color.rgb = color
    if not mapa:
        run_none = paragraph.add_run("Ninguna."); run_none.font.name = font_name; run_none.font.size = font_size
        if color: run_none.font.color.rgb = color
        return
    current_p = paragraph
    for examen, recomendaciones in mapa.items():
        new_p_element = OxmlElement('w:p'); current_p._p.addnext(new_p_element)
        new_para = Paragraph(new_p_element, parent_container)
        new_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
        new_para.paragraph_format.space_after = Pt(2); new_para.paragraph_format.space_before = Pt(3)
        run_examen = new_para.add_run(f"{examen}:")
        run_examen.bold = True; run_examen.font.name = font_name; run_examen.font.size = font_size
        if color: run_examen.font.color.rgb = color
        current_p = new_para
        if not recomendaciones:
            sin_rec_element = OxmlElement('w:p'); current_p._p.addnext(sin_rec_element)
            sin_rec_para = Paragraph(sin_rec_element, parent_container)
            sin_rec_para.paragraph_format.left_indent = Inches(0.28)
            sin_rec_para.paragraph_format.space_after = Pt(2)
            run_sin = sin_rec_para.add_run("Sin recomendación específica registrada en el certificado.")
            run_sin.italic = True; run_sin.font.name = font_name; run_sin.font.size = font_size
            if color: run_sin.font.color.rgb = color
            current_p = sin_rec_para
            continue
        for item in recomendaciones:
            rec_element = OxmlElement('w:p'); current_p._p.addnext(rec_element)
            rec_para = Paragraph(rec_element, parent_container)
            rec_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
            rec_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
            rec_para.paragraph_format.space_after = Pt(2); rec_para.paragraph_format.space_before = Pt(0)
            rec_para.paragraph_format.left_indent = Inches(0.38)
            rec_para.paragraph_format.first_line_indent = Inches(-0.18)
            run_rec = rec_para.add_run("• " + item); run_rec.font.name = font_name; run_rec.font.size = font_size
            if color: run_rec.font.color.rgb = color
            current_p = rec_para

def replace_label_placeholder(paragraph, label_text, placeholder, value):
    if placeholder not in paragraph.text: return False
    if paragraph.runs: font_name, font_size, color = paragraph.runs[0].font.name or "Arial", paragraph.runs[0].font.size or Pt(11), paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color else None
    else: font_name, font_size, color = "Arial", Pt(11), None
    paragraph.text = ""; paragraph.paragraph_format.space_after = Pt(3); paragraph.paragraph_format.space_before = Pt(0)
    run_lbl = paragraph.add_run(label_text); run_lbl.bold = True; run_lbl.font.name = font_name; run_lbl.font.size = font_size
    if color: run_lbl.font.color.rgb = color
    val_clean = value.strip() if value else "Ninguna."
    if es_vacio_o_negativo(val_clean): val_clean = "Ninguna."
    run_val = paragraph.add_run(val_clean); run_val.font.name = font_name; run_val.font.size = font_size
    if color: run_val.font.color.rgb = color
    return True

def obtener_siguiente_consecutivo_local():
    val = obtener_config("ultimo_consecutivo_local")
    return int(val) + 1 if val else 1

def incrementar_consecutivo_local():
    next_num = obtener_siguiente_consecutivo_local()
    guardar_config("ultimo_consecutivo_local", str(next_num))
    return f"SST-2026-{next_num}"

def crear_documento_base():
    """Plantilla funcional de respaldo cuando aún no se ha cargado la institucional."""
    doc = Document()
    seccion = doc.sections[0]
    seccion.top_margin = Inches(0.65)
    seccion.bottom_margin = Inches(0.65)
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("JER S.A.\nRECOMENDACIONES MÉDICAS OCUPACIONALES")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    doc.add_paragraph("Consecutivo: {{NUMERO DE CONSECUTIVO}}")
    doc.add_paragraph("{{LUGAR}}, {{FECHA HOY}}")
    doc.add_paragraph("Señor(a):\n{{NOMBRE DE LA PERSONA}}\n{{CARGO DE LA PERSONA}}")
    doc.add_paragraph("ASUNTO: RECOMENDACIONES EXAMEN {{TIPO DE EXAMEN}}")
    doc.add_paragraph(
        "Según los lineamientos del programa de medicina preventiva y del trabajo de JER S.A; "
        "se hace entrega de las recomendaciones establecidas por el Proveedor de servicios de "
        "Exámenes Médico Ocupacionales (Ingreso, Periódico, egreso, cambio de cargo y post incapacidad)"
    )
    doc.add_paragraph("EXÁMENES REALIZADOS:")
    doc.add_paragraph("{{LISTA DE EXAMENES REALIZADOS}}")
    doc.add_paragraph("{{Recomendaciones médicas}}")
    doc.add_paragraph("Programa de vigilancia epidemiológica: {{Programa de vigilancia epidemiológica}}")
    doc.add_paragraph("{{Observaciones}}")
    doc.add_paragraph("{{Remisiones}}")
    doc.add_paragraph("\nVÍCTOR ALONSO MORENO CASAS\nCoordinador SST")
    return doc

def insertar_firma_en_contenedor(container, firma_file):
    if not firma_file:
        return False
    for idx, parrafo in enumerate(list(container.paragraphs)):
        if "VÍCTOR ALONSO MORENO CASAS" not in parrafo.text.upper():
            continue
        if idx > 0 and not container.paragraphs[idx - 1].text.strip():
            destino = container.paragraphs[idx - 1]
        else:
            nuevo_elemento = OxmlElement("w:p")
            parrafo._p.addprevious(nuevo_elemento)
            destino = Paragraph(nuevo_elemento, container)
        destino.alignment = WD_ALIGN_PARAGRAPH.LEFT
        buffer_firma = archivo_a_buffer(firma_file)
        destino.add_run().add_picture(buffer_firma, width=Inches(1.6))
        return True
    return False

def generar_word_unico(datos_trabajador, lugar, fecha, template_uploaded, firma_file):
    datos_trabajador.update(normalizar_datos_documento(datos_trabajador))
    if template_uploaded: doc_word = Document(archivo_a_buffer(template_uploaded))
    elif os.path.exists("FORMATO RECOMENDACIONES MEDICAS BOT.docx"): doc_word = Document("FORMATO RECOMENDACIONES MEDICAS BOT.docx")
    else: doc_word = crear_documento_base()
    
    consecutivo_final = datos_trabajador.get("consecutivo", "")
    if not consecutivo_final:
        g_url = obtener_config("google_sheets_url")
        if g_url:
            try:
                r = requests.get(g_url, params={"name": datos_trabajador["nombre"], "cargo": datos_trabajador["cargo"], "examen": datos_trabajador["tipo_examen"], "fecha": fecha.strftime("%Y-%m-%d")}, timeout=12)
                r.raise_for_status()
                respuesta_sheets = r.json()
                consecutivo_final = respuesta_sheets.get("consecutive") if respuesta_sheets.get("status") == "success" else incrementar_consecutivo_local()
            except (requests.RequestException, ValueError, TypeError):
                consecutivo_final = incrementar_consecutivo_local()
        else: consecutivo_final = incrementar_consecutivo_local()
        datos_trabajador["consecutivo"] = consecutivo_final

    simple_replacements = {
        "{{NUMERO DE CONSECUTIVO}}": consecutivo_final, "{{TIPO DE EXAMEN}}": datos_trabajador["tipo_examen"],
        "{{LUGAR}}": lugar, "{{FECHA HOY}}": formatear_fecha_es(fecha),
        "{{NOMBRE DE LA PERSONA}}": datos_trabajador["nombre"], "{{CARGO DE LA PERSONA}}": datos_trabajador["cargo"],
        "{{Programa de vigilancia epidemiológica}}": datos_trabajador.get("vigilancia_programa", "Ninguno")
    }

    def procesar_parrafo(p, container):
        if "{{LISTA DE EXAMENES REALIZADOS}}" in p.text:
            insert_bullets_in_placeholder(container, p, datos_trabajador["examenes_lista"])
            return True
        if "{{Recomendaciones médicas}}" in p.text:
            insert_recommendations_in_placeholder(container, p, datos_trabajador)
            return True
        if "{{Observaciones}}".lower() in p.text.lower():
            replace_label_placeholder(p, "Observaciones: ", p.text, datos_trabajador["observaciones"])
            return True
        if "{{Remisiones}}".lower() in p.text.lower():
            replace_label_placeholder(p, "Remisiones: ", p.text, datos_trabajador["remisiones"])
            return True
            
        for k, v in simple_replacements.items():
            if k in p.text: replace_placeholder_in_paragraph_runs(p, k, v)
        aplicar_negrita_dinamica_cuerpo(p, datos_trabajador["tipo_examen"])

    for p in list(doc_word.paragraphs): procesar_parrafo(p, doc_word)
    firma_insertada = insertar_firma_en_contenedor(doc_word, firma_file)
    for table in doc_word.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs): procesar_parrafo(p, cell)
                if not firma_insertada:
                    firma_insertada = insertar_firma_en_contenedor(cell, firma_file)
    b_io = io.BytesIO(); doc_word.save(b_io)
    return b_io.getvalue(), consecutivo_final

def convertir_docx_a_pdf(docx_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(docx_bytes); temp_docx_path = temp_docx.name
    pdf_path = temp_docx_path.replace(".docx", ".pdf")
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(temp_docx_path), temp_docx_path],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=45, check=False
        )
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f: pdf_bytes = f.read()
            os.unlink(temp_docx_path); os.unlink(pdf_path)
            return pdf_bytes, True
    except (OSError, subprocess.SubprocessError):
        pass
    try:
        from docx2pdf import convert
        convert(temp_docx_path, pdf_path)
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f: pdf_bytes = f.read()
            os.unlink(temp_docx_path); os.unlink(pdf_path)
            return pdf_bytes, True
    except Exception:
        pass
    if os.path.exists(temp_docx_path): os.unlink(temp_docx_path)
    return None, False

def formatear_fecha_es(fecha):
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto",
             "septiembre", "octubre", "noviembre", "diciembre"]
    return f"{fecha.day} de {meses[fecha.month - 1]} de {fecha.year}"

def preparar_archivo_final(datos, lugar, fecha, plantilla, firma, formato_salida):
    bytes_word, consecutivo = generar_word_unico(datos, lugar, fecha, plantilla, firma)
    bytes_pdf, pdf_ok = convertir_docx_a_pdf(bytes_word)
    nombre_seguro = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9_-]+", "_", datos.get("nombre", "Trabajador")).strip("_")
    html_vista = generar_html_vista(datos, consecutivo, lugar, fecha)
    if "Word" in formato_salida:
        archivo_bytes = bytes_word
        extension = "docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif "PDF" in formato_salida:
        if not pdf_ok:
            raise RuntimeError("LibreOffice no pudo convertir el Word a PDF.")
        archivo_bytes = bytes_pdf
        extension = "pdf"
        mime = "application/pdf"
    else:
        archivo_bytes = html_vista.encode("utf-8")
        extension = "html"
        mime = "text/html"
    return {
        "bytes": archivo_bytes, "word_bytes": bytes_word, "preview_pdf": bytes_pdf if pdf_ok else None,
        "preview_html": html_vista, "consec_num": consecutivo,
        "filename": f"Recomendaciones_{nombre_seguro}.{extension}", "mime": mime,
        "format": formato_salida
    }

def _bytes_recurso_generacion(recurso):
    if not recurso:
        return b""
    if isinstance(recurso, dict):
        return bytes(recurso.get("bytes", b"") or b"")
    if isinstance(recurso, (bytes, bytearray)):
        return bytes(recurso)
    if hasattr(recurso, "getvalue"):
        return recurso.getvalue()
    return b""

def calcular_huella_salida(datos, plantilla, firma, formato_salida):
    """Detecta cambios reales sin volver a extraer ni validar los PDF."""
    datos_salida = {
        clave: datos.get(clave) for clave in sorted(CAMPOS_QUE_AFECTAN_SALIDA)
    }
    payload = {
        "datos": datos_salida,
        "formato": formato_salida,
        "plantilla": hashlib.sha256(_bytes_recurso_generacion(plantilla)).hexdigest(),
        "firma": hashlib.sha256(_bytes_recurso_generacion(firma)).hexdigest()
    }
    serializado = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
    return hashlib.sha256(serializado.encode("utf-8")).hexdigest()

def construir_zip_salidas(salidas):
    if not salidas:
        return None
    zip_buffer = io.BytesIO()
    nombres_usados = set()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for indice, archivo in enumerate(salidas.values(), start=1):
            nombre_zip = archivo["filename"]
            if nombre_zip in nombres_usados:
                base, extension = os.path.splitext(nombre_zip)
                nombre_zip = f"{base}_{indice}{extension}"
            nombres_usados.add(nombre_zip)
            zf.writestr(nombre_zip, archivo["bytes"])
    return zip_buffer.getvalue()

def obtener_o_generar_salida(
    nombre_pdf, datos, plantilla, firma, formato_salida,
    cache_salidas=None, cache_huellas=None
):
    cache_salidas = cache_salidas or {}
    cache_huellas = cache_huellas or {}
    huella_actual = calcular_huella_salida(datos, plantilla, firma, formato_salida)
    archivo_cache = cache_salidas.get(nombre_pdf)
    if archivo_cache and cache_huellas.get(nombre_pdf) == huella_actual:
        return archivo_cache, huella_actual, True
    archivo = preparar_archivo_final(
        datos, datos.get("lugar", "Tunja"), datos.get("fecha", datetime.date.today()),
        plantilla, firma, formato_salida
    )
    return archivo, huella_actual, False

def preparar_lote_documentos(
    documentos, plantilla, firma, formato_salida, progreso_callback=None,
    cache_salidas=None, cache_huellas=None
):
    """Generación incremental: reutiliza cada documento vigente y procesa solo faltantes o editados."""
    if not documentos:
        raise ValueError("No hay documentos cargados para generar.")
    cache_salidas = cache_salidas or {}
    cache_huellas = cache_huellas or {}
    resultados = {}
    huellas = {}
    generados = 0
    reutilizados = 0
    total = len(documentos)
    for indice, (nombre_pdf, datos_trabajador) in enumerate(documentos.items(), start=1):
        archivo, huella, reutilizado = obtener_o_generar_salida(
            nombre_pdf, datos_trabajador, plantilla, firma, formato_salida,
            cache_salidas, cache_huellas
        )
        resultados[nombre_pdf] = archivo
        huellas[nombre_pdf] = huella
        if reutilizado:
            reutilizados += 1
        else:
            generados += 1
        if progreso_callback:
            progreso_callback(indice, total, nombre_pdf, datos_trabajador, reutilizado)
    return resultados, construir_zip_salidas(resultados), huellas, generados, reutilizados

def correo_valido(correo):
    correo = str(correo or "").strip()
    nombre, direccion = parseaddr(correo)
    return bool(direccion and direccion == correo and re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", direccion))

def separar_correos(correos):
    """Acepta uno o varios correos separados por coma, punto y coma o salto de línea."""
    resultado = []
    vistos = set()
    if isinstance(correos, (list, tuple, set)):
        elementos = []
        for grupo in correos:
            elementos.extend(re.split(r"[,;\n]+", str(grupo or "")))
    else:
        elementos = re.split(r"[,;\n]+", str(correos or ""))
    for correo in elementos:
        correo_limpio = correo.strip().lower()
        if correo_limpio and correo_limpio not in vistos:
            resultado.append(correo_limpio)
            vistos.add(correo_limpio)
    return resultado

def correos_invalidos(correos):
    return [correo for correo in separar_correos(correos) if not correo_valido(correo)]

def obtener_configuracion_smtp():
    host = obtener_config("smtp_host").strip()
    password = obtener_config("smtp_password")
    if host.lower() in {"smtp.gmail.com", "smtp.googlemail.com"}:
        password = re.sub(r"\s+", "", password)
    return {
        "host": host,
        "port": int(obtener_config("smtp_port") or 587),
        "user": obtener_config("smtp_user").strip(),
        "password": password,
        "sender": (obtener_config("smtp_sender") or obtener_config("smtp_user")).strip(),
        "use_tls": obtener_config("smtp_use_tls") != "0"
    }

def personalizar_plantilla_correo(plantilla, datos):
    nombre = datos.get("nombre", "").strip() or "el colaborador"
    identificacion = datos.get("identificacion", "").strip() or "no registrada"
    return (
        str(plantilla or "")
        .replace("{nombre}", nombre)
        .replace("{identificacion}", identificacion)
    )

def generar_mensaje_correo(datos, plantilla=None):
    return personalizar_plantilla_correo(
        plantilla if plantilla is not None else PLANTILLA_CUERPO_CORREO, datos
    )

def generar_asunto_correo(datos, plantilla=None):
    return personalizar_plantilla_correo(
        plantilla if plantilla is not None else PLANTILLA_ASUNTO_CORREO, datos
    )

def enviar_archivo_por_correo(
    destinatario, asunto, mensaje, archivo, configuracion, cc=None, cco=None
):
    """Envía un único documento; la interfaz exige selección y confirmación explícitas."""
    if not correo_valido(destinatario):
        raise ValueError(f"El correo {destinatario or '(vacío)'} no es válido.")
    lista_cc = separar_correos(cc)
    lista_cco = separar_correos(cco)
    adicionales_invalidos = [
        correo for correo in lista_cc + lista_cco if not correo_valido(correo)
    ]
    if adicionales_invalidos:
        raise ValueError(
            "Revisa los correos de copia: " + ", ".join(adicionales_invalidos)
        )
    destino_normalizado = destinatario.strip().lower()
    lista_cc = [correo for correo in lista_cc if correo != destino_normalizado]
    lista_cco = [
        correo for correo in lista_cco
        if correo != destino_normalizado and correo not in lista_cc
    ]
    if not configuracion.get("host") or not configuracion.get("sender"):
        raise ValueError("Configura el servidor SMTP y el correo remitente.")
    email = EmailMessage()
    email["From"] = configuracion["sender"]
    email["To"] = destinatario
    if lista_cc:
        email["Cc"] = ", ".join(lista_cc)
    if lista_cco:
        email["Bcc"] = ", ".join(lista_cco)
    email["Subject"] = asunto
    email.set_content(mensaje)
    tipo_mime, subtipo_mime = archivo["mime"].split("/", 1)
    email.add_attachment(
        archivo["bytes"], maintype=tipo_mime, subtype=subtipo_mime,
        filename=archivo["filename"]
    )
    contexto = ssl.create_default_context()
    puerto = int(configuracion.get("port") or 587)
    if puerto == 465:
        with smtplib.SMTP_SSL(configuracion["host"], puerto, context=contexto, timeout=30) as servidor:
            if configuracion.get("user"):
                servidor.login(configuracion["user"], configuracion.get("password", ""))
            servidor.send_message(email)
    else:
        with smtplib.SMTP(configuracion["host"], puerto, timeout=30) as servidor:
            servidor.ehlo()
            if configuracion.get("use_tls", True):
                servidor.starttls(context=contexto)
                servidor.ehlo()
            if configuracion.get("user"):
                servidor.login(configuracion["user"], configuracion.get("password", ""))
            servidor.send_message(email)
    return True

def mostrar_previsualizacion_pdf(pdf_bytes):
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as preview_pdf:
        for i, page in enumerate(preview_pdf.pages):
            imagen = page.to_image(resolution=145).original
            st.image(imagen, caption=f"Vista final · Página {i + 1}", use_container_width=True)

def generar_html_vista(datos, consecutivo_num, lugar, fecha):
    seguro = lambda valor: html.escape(str(valor or ""))
    examenes_html = "".join(f"<li>{seguro(ex)}</li>" for ex in datos.get("examenes_lista", []))
    mapa_recomendaciones = agrupar_recomendaciones_por_examen(
        datos.get("examenes_lista", []), datos.get("recomendaciones_lista", []),
        datos.get("recomendaciones_por_examen", {})
    )
    bloques_recomendaciones = []
    for examen, recomendaciones in mapa_recomendaciones.items():
        if recomendaciones:
            lista = "".join(f"<li>{seguro(rec)}</li>" for rec in recomendaciones)
        else:
            lista = "<li><em>Sin recomendación específica registrada en el certificado.</em></li>"
        bloques_recomendaciones.append(
            f"<div style='margin:10px 0 4px'><strong>{seguro(examen)}:</strong></div><ul style='margin-top:4px'>{lista}</ul>"
        )
    recomendaciones_html = "".join(bloques_recomendaciones) or "<p>Ninguna.</p>"
    return f"""
    <div style="font-family: Arial, sans-serif; color: #333; padding: 20px; line-height: 1.5; background: white; border: 1px solid #ccc; max-width: 800px; margin: auto;">
        <div style="text-align: right; font-weight: bold; color: #1f4e79;">Consecutivo: {seguro(consecutivo_num)}</div>
        <div style="text-align: center; font-weight: bold; font-size: 16px; margin: 20px 0; color: #1f4e79; background: #f0f4f8; padding: 8px;">
            ASUNTO: RECOMENDACIONES EXAMEN {seguro(datos.get('tipo_examen', ''))}
        </div>
        <div>{seguro(lugar)}, {formatear_fecha_es(fecha)}</div><br>
        <div>Sr(a).<br><strong>{seguro(datos.get('nombre'))}</strong><br>{seguro(datos.get('cargo'))}</div><br>
        <p>Cordial saludo,</p>
        <p>Según los lineamientos del programa de medicina preventiva y del trabajo de JER S.A; se hace entrega de las recomendaciones establecidas por el Proveedor de servicios de Exámenes Médico Ocupacionales (Ingreso, Periódico, egreso, cambio de cargo y post incapacidad)</p>
        <p><strong>EXÁMENES REALIZADOS:</strong></p>
        <ul>{examenes_html}</ul>
        <p><strong>Recomendaciones por examen:</strong></p>
        {recomendaciones_html}
        <p><strong>Programa de Vigilancia:</strong> {seguro(datos.get('vigilancia_programa', 'NINGUNO'))}</p>
        <p><strong>Observaciones:</strong> {seguro(datos.get('observaciones'))}</p>
        <p><strong>Remisiones:</strong> {seguro(datos.get('remisiones'))}</p><br>
        <p>Atentamente,</p><br>
        <p><strong>VÍCTOR ALONSO MORENO CASAS</strong><br>Coordinador SST</p>
    </div>
    """

# --- GESTIÓN DE ACCESOS Y AUTENTICACIÓN ---
if not st.session_state.logged_in:
    st.markdown("<div class='login-box'>", unsafe_allow_html=True)
    st.markdown("<h2>🔑 Acceso Seguro</h2>", unsafe_allow_html=True)
    st.markdown("<p>Portal Interno de Medicina Preventiva - JER S.A.</p>", unsafe_allow_html=True)
    if not tiene_usuarios():
        st.warning("🆕 Bienvenido. Configura tu cuenta inicial de Administrador.")
        with st.form("form_registro_inicial"):
            reg_nombre = st.text_input("Nombre Completo", key="init_admin_fullname")
            reg_user = st.text_input("Nombre de Usuario (Login)", key="init_admin_username")
            reg_pwd = st.text_input("Contraseña", type="password", key="init_admin_password")
            if st.form_submit_button("Crear Administrador"):
                if reg_nombre and reg_user and reg_pwd:
                    if registrar_usuario(reg_user, reg_pwd, reg_nombre): st.success("¡Administrador creado con éxito!"); st.rerun()
                else: st.warning("Completa todos los campos.")
    else:
        opcion_acceso = st.radio("Elige una acción:", ["Iniciar Sesión", "Crear Nueva Cuenta", "Actualizar Contraseña"], horizontal=True, key="sistema_tabs_acceso")
        st.markdown("<br>", unsafe_allow_html=True)
        if opcion_acceso == "Iniciar Sesión":
            with st.form("form_inicio_sesion"):
                log_user = st.text_input("Usuario", key="login_username_field")
                log_pwd = st.text_input("Contraseña", type="password", key="login_password_field")
                if st.form_submit_button("Ingresar al Sistema"):
                    nombre_usuario = verificar_usuario(log_user, log_pwd)
                    if nombre_usuario: st.session_state.logged_in = True; st.session_state.username = nombre_usuario; st.rerun()
                    else: st.error("❌ Credenciales incorrectas.")
        elif opcion_acceso == "Crear Nueva Cuenta":
            with st.form("form_crear_cuenta"):
                reg_nombre = st.text_input("Nombre Completo", key="register_fullname_field")
                reg_user = st.text_input("Nombre de Usuario", key="register_username_field")
                reg_pwd = st.text_input("Contraseña", type="password", key="register_password_field")
                if st.form_submit_button("Registrar Cuenta"):
                    if reg_nombre and reg_user and reg_pwd:
                        if registrar_usuario(reg_user, reg_pwd, reg_nombre): st.success("🎉 Cuenta creada. Inicia Sesión.")
                        else: st.error("❌ El usuario ya existe.")
                    else: st.warning("Completa todos los campos.")
        elif opcion_acceso == "Actualizar Contraseña":
            with st.form("form_update_password"):
                upd_user = st.text_input("Usuario", key="update_username_field")
                upd_old_pwd = st.text_input("Contraseña Actual", type="password", key="update_old_password_field")
                upd_new_pwd = st.text_input("Nueva Contraseña", type="password", key="update_new_password_field")
                if st.form_submit_button("Cambiar Contraseña"):
                    if upd_user and upd_old_pwd and upd_new_pwd:
                        if actualizar_contrasena(upd_user, upd_old_pwd, upd_new_pwd): st.success("✅ Contraseña actualizada.")
                        else: st.error("❌ Error en los datos proporcionados.")
    st.markdown("</div>", unsafe_allow_html=True); st.stop()

# --- BANNER PRINCIPAL Y CONFIGURACIÓN PERSISTENTE ---
st.markdown(
    "<div class='header-banner'><h1>🩺 Comunicaciones Médicas SST</h1>"
    "<p>Extracción verificable, edición controlada y previsualización antes de descargar.</p></div>",
    unsafe_allow_html=True
)

st.sidebar.markdown("### 👤 Sesión activa")
st.sidebar.markdown(f"<div class='metric-card'><strong>{html.escape(st.session_state.username)}</strong><br>Portal interno JER S.A.</div>", unsafe_allow_html=True)
st.sidebar.caption(f"Aplicación: {APP_VERSION} · Motor: {PROCESSING_PIPELINE_VERSION}")
if st.sidebar.button("Cerrar sesión", key="logout_button"):
    for clave in ["documentos", "pdfs_raw_bytes", "textos_raw", "ai_validation"]:
        st.session_state[clave] = {}
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.session_state.processed_doc = None
    st.session_state.zip_bytes = None
    st.session_state.batch_outputs = {}
    st.session_state.output_fingerprints = {}
    st.session_state.last_cache_stats = {"generados": 0, "reutilizados": 0}
    st.session_state.email_send_feedback = None
    st.session_state.prev_colaborador = None
    st.session_state.document_count = 0
    st.rerun()

with st.sidebar.expander("🔗 Google Sheets", expanded=True):
    g_url_guardada = obtener_config("google_sheets_url")
    g_url_input = st.text_input("URL de Google Apps Script", value=g_url_guardada, type="password", key="google_sheets_url_input")
    if st.button("Guardar conexión", key="save_google_connection"):
        guardar_config("google_sheets_url", g_url_input.strip())
        g_url_guardada = g_url_input.strip()
        st.success("Conexión guardada de forma persistente.")
    st.caption("🟢 Configurada" if g_url_guardada else "🟡 Se usará consecutivo local")

with st.sidebar.expander("📄 Plantilla y firma", expanded=True):
    nueva_plantilla = st.file_uploader("Plantilla institucional (.docx)", type=["docx"], key="persistent_template_upload")
    nueva_firma = st.file_uploader("Firma autorizada (.png, .jpg)", type=["png", "jpg", "jpeg"], key="persistent_signature_upload")
    for clave, archivo in [("plantilla_word", nueva_plantilla), ("firma_autorizada", nueva_firma)]:
        if archivo is not None:
            digest = hashlib.sha256(archivo.getvalue()).hexdigest()
            if st.session_state.asset_hashes.get(clave) != digest:
                guardar_archivo_config(clave, archivo)
                st.session_state.asset_hashes[clave] = digest
    plantilla_guardada = obtener_archivo_config("plantilla_word")
    firma_guardada = obtener_archivo_config("firma_autorizada")
    st.caption(f"✅ Plantilla: {plantilla_guardada['name']}" if plantilla_guardada else "⚠️ Sin plantilla; se usará el formato de respaldo")
    st.caption(f"✅ Firma: {firma_guardada['name']}" if firma_guardada else "⚠️ Sin firma guardada")
    col_borrar_1, col_borrar_2 = st.columns(2)
    with col_borrar_1:
        if st.button("Quitar plantilla", disabled=not plantilla_guardada, key="remove_template"):
            eliminar_archivo_config("plantilla_word")
            st.rerun()
    with col_borrar_2:
        if st.button("Quitar firma", disabled=not firma_guardada, key="remove_signature"):
            eliminar_archivo_config("firma_autorizada")
            st.rerun()

with st.sidebar.expander("✨ Validación con IA", expanded=False):
    clave_ia_nueva = st.text_input("Clave API de Gemini", type="password", placeholder="Déjala vacía si ya está guardada", key="gemini_key_input")
    modelo_guardado = obtener_config("gemini_model") or GEMINI_MODEL_DEFAULT
    if modelo_guardado in {"gemini-2.5-flash", "models/gemini-2.5-flash"}:
        modelo_guardado = GEMINI_MODEL_DEFAULT
        guardar_config("gemini_model", modelo_guardado)
    modelo_ia = st.text_input("Modelo", value=modelo_guardado, key="gemini_model_input")
    if st.button("Guardar configuración de IA", key="save_ai_configuration"):
        if clave_ia_nueva.strip():
            guardar_config("gemini_api_key", clave_ia_nueva.strip())
        guardar_config("gemini_model", modelo_ia.strip() or GEMINI_MODEL_DEFAULT)
        st.success("Configuración de IA guardada.")
    st.caption("🟢 Clave configurada" if obtener_gemini_api_key() else "🟡 IA opcional sin configurar")

with st.sidebar.expander("✉️ Correo saliente", expanded=False):
    smtp_host_input = st.text_input("Servidor SMTP", value=obtener_config("smtp_host"), placeholder="smtp.empresa.com", key="smtp_host_input")
    smtp_port_actual = int(obtener_config("smtp_port") or 587)
    smtp_port_input = st.number_input("Puerto", min_value=1, max_value=65535, value=smtp_port_actual, step=1, key="smtp_port_input")
    smtp_user_input = st.text_input("Usuario SMTP", value=obtener_config("smtp_user"), key="smtp_user_input")
    smtp_sender_input = st.text_input("Correo remitente", value=obtener_config("smtp_sender") or obtener_config("smtp_user"), key="smtp_sender_input")
    smtp_password_input = st.text_input("Contraseña o clave de aplicación", type="password", placeholder="Déjala vacía para conservar la guardada", key="smtp_password_input")
    smtp_tls_input = st.checkbox("Usar conexión segura TLS", value=obtener_config("smtp_use_tls") != "0", key="smtp_tls_input")
    if st.button("Guardar configuración de correo", key="save_smtp_configuration"):
        guardar_config("smtp_host", smtp_host_input.strip())
        guardar_config("smtp_port", str(int(smtp_port_input)))
        guardar_config("smtp_user", smtp_user_input.strip())
        guardar_config("smtp_sender", smtp_sender_input.strip())
        guardar_config("smtp_use_tls", "1" if smtp_tls_input else "0")
        if smtp_password_input:
            password_guardada = smtp_password_input
            if smtp_host_input.strip().lower() in {"smtp.gmail.com", "smtp.googlemail.com"}:
                password_guardada = re.sub(r"\s+", "", password_guardada)
            guardar_config("smtp_password", password_guardada)
        st.success("Configuración de correo guardada.")
    st.markdown(
        "[🔑 Crear una nueva contraseña de aplicación de Google]"
        "(https://myaccount.google.com/apppasswords)"
    )
    st.caption(
        "Google no permite volver a consultar una clave anterior. Si la perdiste, "
        "crea una nueva, pégala en el campo de contraseña y guarda la configuración."
    )
    st.caption("Los mensajes nunca se envían automáticamente: siempre requieren selección y confirmación.")

plantilla_activa = plantilla_guardada
firma_activa = firma_guardada
api_key_ia = obtener_gemini_api_key()
modelo_ia = obtener_config("gemini_model") or GEMINI_MODEL_DEFAULT
configuracion_smtp = obtener_configuracion_smtp()

st.markdown(f"""
<div class="status-row">
  <div class="status-card"><strong>Google Sheets</strong><span class="{'status-ok' if g_url_guardada else 'status-warn'}">{'Conectado y guardado' if g_url_guardada else 'Consecutivo local'}</span></div>
  <div class="status-card"><strong>Documentación base</strong><span class="{'status-ok' if plantilla_activa and firma_activa else 'status-warn'}">{'Plantilla y firma listas' if plantilla_activa and firma_activa else 'Revisa plantilla o firma'}</span></div>
  <div class="status-card"><strong>Motor de validación</strong><span class="{'status-ok' if api_key_ia else 'status-warn'}">{'Gemini disponible' if api_key_ia else 'Extractor local activo'}</span></div>
  <div class="status-card"><strong>Versión activa</strong><span class="status-ok">{APP_VERSION}</span></div>
</div>
""", unsafe_allow_html=True)

st.markdown("<div class='section-title'><span class='section-number'>1</span>Cargar certificados médicos</div>", unsafe_allow_html=True)
pdfs_subidos = st.file_uploader(
    "Selecciona uno o varios certificados en PDF",
    type="pdf", accept_multiple_files=True, key="medical_pdf_uploader",
    help="La extracción local revisa texto, tablas, coordenadas y OCR cuando está disponible."
)

archivo_seleccionado = None
if pdfs_subidos:
    firma_lote = (
        PROCESSING_PIPELINE_VERSION,
        tuple((pdf.name, hashlib.sha256(pdf.getvalue()).hexdigest()) for pdf in pdfs_subidos)
    )
    if st.session_state.get("document_signature") != firma_lote:
        st.session_state.documentos = {}
        st.session_state.pdfs_raw_bytes = {}
        st.session_state.textos_raw = {}
        st.session_state.processed_doc = None
        st.session_state.zip_bytes = None
        st.session_state.batch_outputs = {}
        st.session_state.output_fingerprints = {}
        st.session_state.last_cache_stats = {"generados": 0, "reutilizados": 0}
        st.session_state.original_batch_preview = {}
        st.session_state.document_signature = firma_lote
        st.session_state.document_count = len(pdfs_subidos)

    pendientes = [pdf for pdf in pdfs_subidos if pdf.name not in st.session_state.documentos]
    if pendientes:
        progreso = st.progress(0, text="Iniciando validación automática de los PDF...")
        for indice, pdf in enumerate(pendientes, start=1):
            pdf_raw_data = pdf.getvalue()
            texto_raw = extraer_texto_pdf_robusto(pdf_raw_data)
            metadatos_pdf = extraer_metadatos_pdf_estructurados(pdf_raw_data, texto_raw)
            st.session_state.pdfs_raw_bytes[pdf.name] = pdf_raw_data
            st.session_state.textos_raw[pdf.name] = texto_raw
            datos_locales = analizar_pdf_inteligente(texto_raw, metadatos_pdf=metadatos_pdf)
            progreso.progress(
                (indice - 0.5) / len(pendientes),
                text=f"Motor local listo. Validando automáticamente con IA: {pdf.name}"
            )
            if api_key_ia:
                try:
                    datos_ia = validar_documento_con_gemini(
                        pdf_raw_data, texto_raw, datos_locales, api_key_ia, modelo_ia
                    )
                    modelo_usado = datos_ia.get("_modelo_usado", GEMINI_MODEL_DEFAULT)
                    guardar_config("gemini_model", modelo_usado)
                    datos_finales = fusionar_validacion_ia(datos_locales, datos_ia, texto_raw)
                    datos_finales["error_ia"] = ""
                    st.session_state.ai_validation[pdf.name] = {
                        "status": "ok", "model": modelo_usado,
                        "second_review": bool(datos_ia.get("_segunda_revision_ia")),
                        "pending_fragments": len(datos_ia.get("_fragmentos_pendientes", []))
                    }
                except (requests.RequestException, RuntimeError, ValueError, json.JSONDecodeError) as exc:
                    datos_finales = datos_locales
                    datos_finales["validado_ia"] = False
                    datos_finales["modo_validacion"] = "Respaldo local sin conexión IA"
                    datos_finales["error_ia"] = "No hubo conexión con la IA; se conservó la extracción local."
                    st.session_state.ai_validation[pdf.name] = {"status": "fallback", "error": str(exc)}
            else:
                datos_finales = datos_locales
                datos_finales["validado_ia"] = False
                datos_finales["modo_validacion"] = "Respaldo local sin clave IA"
                datos_finales["error_ia"] = "Configura la clave de Gemini para habilitar la validación automática."
                st.session_state.ai_validation[pdf.name] = {"status": "local"}
            st.session_state.documentos[pdf.name] = normalizar_datos_documento(datos_finales)
            progreso.progress(indice / len(pendientes), text=f"Completado: {pdf.name}")
        progreso.empty()

    col_m1, col_m2, col_m3 = st.columns(3)
    col_m1.metric("Certificados", len(st.session_state.documentos))
    col_m2.metric("Exámenes detectados", sum(len(d.get("examenes_lista", [])) for d in st.session_state.documentos.values()))
    col_m3.metric("Validados con IA", sum(bool(d.get("validado_ia")) for d in st.session_state.documentos.values()))
    st.caption(
        "El detalle técnico de cada archivo está disponible en el apartado "
        "‘Control y tablas’, para mantener esta pantalla más sencilla."
    )
    if api_key_ia and all(d.get("validado_ia") for d in st.session_state.documentos.values()):
        st.success("Todos los PDF del lote fueron validados automáticamente con IA al cargarse.")
    elif api_key_ia:
        st.warning("Algunos PDF usaron el respaldo local porque la IA o la conexión no respondieron.")
    else:
        st.warning("Los PDF se procesaron con el respaldo local. Configura Gemini para la validación automática con IA.")
    archivo_seleccionado = st.selectbox("Selecciona el certificado que deseas revisar", list(st.session_state.documentos.keys()), key="selected_worker_pdf")
else:
    if st.session_state.documentos:
        st.session_state.documentos = {}
        st.session_state.pdfs_raw_bytes = {}
        st.session_state.textos_raw = {}
        st.session_state.processed_doc = None
        st.session_state.batch_outputs = {}
        st.session_state.output_fingerprints = {}
        st.session_state.last_cache_stats = {"generados": 0, "reutilizados": 0}
        st.session_state.zip_bytes = None
        st.session_state.original_batch_preview = {}
        st.session_state.document_signature = None

if not archivo_seleccionado:
    st.info("Carga al menos un PDF para activar el editor, la validación y la previsualización.")
    st.stop()

if archivo_seleccionado != st.session_state.prev_colaborador:
    st.session_state.processed_doc = st.session_state.batch_outputs.get(archivo_seleccionado)
    st.session_state.prev_colaborador = archivo_seleccionado
    st.session_state.editor_version += 1

doc_actual = st.session_state.documentos[archivo_seleccionado]
id_editor = hashlib.sha1(archivo_seleccionado.encode("utf-8")).hexdigest()[:8]
prefijo = f"{id_editor}_{st.session_state.editor_version}"

st.markdown("<div class='section-title'><span class='section-number'>2</span>Revisar, previsualizar y distribuir</div>", unsafe_allow_html=True)
st.info("Flujo disponible: revisar datos → comprobar originales → generar → enviar. Los reportes tabulares están separados en Control y tablas.")
tab_datos, tab_origen, tab_vista, tab_lote, tab_control = st.tabs([
    "📝 Datos", "🔎 PDF original", "👁️ Vista final", "✉️ Correo", "📊 Control y tablas"
])

resumen_confirmacion = []

incompletos_lote = [
    nombre_pdf for nombre_pdf, datos in st.session_state.documentos.items()
    if (
        not datos.get("nombre") or not datos.get("examenes_lista")
        or datos.get("recomendaciones_pendientes_revision")
    )
]

with tab_datos:
    if doc_actual.get("validado_ia"):
        st.success("Este certificado ya fue contrastado por el motor de IA y por las reglas locales.")
    col_f1, col_f2 = st.columns(2)
    with col_f1:
        lugar = st.text_input("Lugar", value=doc_actual.get("lugar", "Tunja"), key=f"lugar_{prefijo}")
    with col_f2:
        fecha = st.date_input("Fecha", value=doc_actual.get("fecha", datetime.date.today()), key=f"fecha_{prefijo}")
    tipo_examen = st.text_input("Tipo de examen", value=doc_actual.get("tipo_examen", "Periódico"), key=f"tipo_{prefijo}")
    col_p1, col_p2 = st.columns(2)
    with col_p1:
        nombre_persona = st.text_input("Trabajador", value=doc_actual.get("nombre", ""), key=f"nombre_{prefijo}")
    with col_p2:
        cargo_persona = st.text_input("Cargo", value=doc_actual.get("cargo", ""), key=f"cargo_{prefijo}")
    col_c1, col_c2 = st.columns(2)
    with col_c1:
        identificacion_persona = st.text_input(
            "Número de identificación", value=doc_actual.get("identificacion", ""), key=f"identificacion_{prefijo}"
        )
    with col_c2:
        correo_persona = st.text_input(
            "Correo extraído del PDF", value=doc_actual.get("correo", ""), key=f"correo_{prefijo}"
        )
    examenes_realizados = st.text_area(
        "Exámenes realizados · uno por línea", value="\n".join(doc_actual.get("examenes_lista", [])),
        height=150, key=f"examenes_{prefijo}"
    )
    examenes_editados = normalizar_lista_clinica(examenes_realizados.splitlines())
    mapa_editor_base = agrupar_recomendaciones_por_examen(
        examenes_editados,
        doc_actual.get("recomendaciones_lista", []),
        doc_actual.get("recomendaciones_por_examen", {})
    )
    st.markdown("#### Recomendaciones organizadas por examen")
    st.caption("Cada campo corresponde a un examen realizado. Escribe una recomendación completa por línea.")
    mapa_editor = {}
    pendientes_editor = []
    for examen in examenes_editados:
        clave_examen = hashlib.sha1(normalizar_etiqueta(examen).encode("utf-8")).hexdigest()[:8]
        texto_examen = st.text_area(
            f"{examen}:",
            value="\n".join(mapa_editor_base.get(examen, [])),
            height=max(100, min(210, 62 + 28 * max(1, len(mapa_editor_base.get(examen, []))))),
            key=f"recomendaciones_{prefijo}_{clave_examen}"
        )
        recomendaciones_examen, pendientes_examen = separar_recomendaciones_atomicas(texto_examen.splitlines())
        mapa_editor[examen] = [
            _separar_prefijo_examen(item)[1] for item in recomendaciones_examen
        ]
        pendientes_editor.extend(
            f"{examen}: {_separar_prefijo_examen(item)[1]}" for item in pendientes_examen
        )
    generales_previas = mapa_editor_base.get("Recomendaciones generales", [])
    recomendaciones_generales = st.text_area(
        "Recomendaciones generales (solo cuando el PDF no las asocia a un examen):",
        value="\n".join(generales_previas), height=110,
        key=f"recomendaciones_generales_{prefijo}"
    )
    generales_editor, generales_pendientes = separar_recomendaciones_atomicas(
        recomendaciones_generales.splitlines()
    )
    if generales_editor:
        mapa_editor["Recomendaciones generales"] = [
            _separar_prefijo_examen(item)[1] for item in generales_editor
        ]
    pendientes_editor.extend(generales_pendientes)
    pendientes_actuales = doc_actual.get("recomendaciones_pendientes_revision", [])
    fragmentos_revision = st.text_area(
        "Fragmentos pendientes de confirmar · complétalos revisando el PDF o elimínalos",
        value="\n".join(pendientes_actuales), height=110,
        key=f"fragmentos_revision_{prefijo}",
        help="Estos fragmentos no pasan al Word mientras sigan evidentemente cortados."
    ) if pendientes_actuales else ""
    programa_vigilancia = st.text_input(
        "Programa de vigilancia epidemiológica (PVE)",
        value=doc_actual.get("vigilancia_programa", "NINGUNO"), key=f"pve_{prefijo}"
    )
    observaciones = st.text_area("Observaciones", value=doc_actual.get("observaciones", ""), height=120, key=f"observaciones_{prefijo}")
    remisiones = st.text_input("Remisiones", value=doc_actual.get("remisiones", "No"), key=f"remisiones_{prefijo}")

    recuperadas_revision, pendientes_revision = separar_recomendaciones_atomicas(
        fragmentos_revision.splitlines() if fragmentos_revision else []
    )
    mapa_editor = agrupar_recomendaciones_por_examen(
        examenes_editados, recuperadas_revision, mapa_editor
    )
    valores_actualizados = {
        "nombre": nombre_persona.strip().title(), "cargo": a_caso_oracion(cargo_persona),
        "identificacion": re.sub(r"\D", "", identificacion_persona), "correo": correo_persona.strip().lower(),
        "tipo_examen": a_caso_oracion(tipo_examen),
        "examenes_lista": examenes_editados,
        "recomendaciones_por_examen": mapa_editor,
        "recomendaciones_lista": aplanar_recomendaciones_por_examen(mapa_editor),
        "recomendaciones_pendientes_revision": deduplicar_textos(
            pendientes_editor + pendientes_revision
        ),
        "observaciones": a_caso_oracion(observaciones), "remisiones": a_caso_oracion(remisiones) or "No",
        "vigilancia_programa": a_caso_oracion(programa_vigilancia), "lugar": lugar.strip().title(), "fecha": fecha
    }
    salida_seleccionada_cambio = False
    for clave, valor in valores_actualizados.items():
        if doc_actual.get(clave) != valor:
            doc_actual[clave] = valor
            if clave in CAMPOS_QUE_AFECTAN_SALIDA:
                salida_seleccionada_cambio = True
    if salida_seleccionada_cambio:
        st.session_state.processed_doc = None
        st.session_state.batch_outputs.pop(archivo_seleccionado, None)
        st.session_state.output_fingerprints.pop(archivo_seleccionado, None)
        st.session_state.zip_bytes = construir_zip_salidas(st.session_state.batch_outputs)

    incompletos_lote = [
        nombre_pdf for nombre_pdf, datos in st.session_state.documentos.items()
        if (
            not datos.get("nombre") or not datos.get("examenes_lista")
            or datos.get("recomendaciones_pendientes_revision")
        )
    ]

    faltantes = []
    if not doc_actual.get("nombre"): faltantes.append("nombre")
    if not doc_actual.get("cargo"): faltantes.append("cargo")
    if not doc_actual.get("identificacion"): faltantes.append("identificación")
    if not doc_actual.get("examenes_lista"): faltantes.append("exámenes realizados")
    if not doc_actual.get("recomendaciones_lista"): faltantes.append("recomendaciones")
    if faltantes:
        st.warning("Revisa estos campos antes de generar: " + ", ".join(faltantes) + ".")
    if doc_actual.get("recomendaciones_pendientes_revision"):
        st.error(
            f"Hay {len(doc_actual['recomendaciones_pendientes_revision'])} fragmento(s) cortado(s). "
            "No se incluirán en el documento hasta que queden completos."
        )

    st.markdown("#### Estado de validación automática")
    if doc_actual.get("validado_ia"):
        detalle_revision = " con segunda revisión automática" if doc_actual.get("segunda_revision_ia") else ""
        st.success(
            f"Validado automáticamente{detalle_revision} con "
            f"{doc_actual.get('modelo_ia', GEMINI_MODEL_DEFAULT)}. "
            "Las recomendaciones quedaron separadas y deduplicadas."
        )
    else:
        st.warning(doc_actual.get("error_ia", "Se utilizó el motor local de respaldo."))

    st.markdown("---")
    formato_salida = st.radio(
        "Formato del archivo final",
        ["Microsoft Word (.docx)", "Documento PDF Oficial (.pdf)", "Impresión de Respaldo Web (HTML)"],
        horizontal=True, key="output_format_global"
    )
    # Comprobar vigencia solo con huellas rápidas; no vuelve a ejecutar OCR ni IA.
    for nombre_cache in list(st.session_state.batch_outputs):
        datos_cache = st.session_state.documentos.get(nombre_cache)
        huella_vigente = (
            calcular_huella_salida(
                datos_cache, plantilla_activa, firma_activa, formato_salida
            ) if datos_cache else None
        )
        if st.session_state.output_fingerprints.get(nombre_cache) != huella_vigente:
            st.session_state.batch_outputs.pop(nombre_cache, None)
            st.session_state.output_fingerprints.pop(nombre_cache, None)
    st.session_state.zip_bytes = construir_zip_salidas(st.session_state.batch_outputs)
    st.session_state.processed_doc = st.session_state.batch_outputs.get(archivo_seleccionado)

    total_documentos = len(st.session_state.documentos)
    total_listos = len(st.session_state.batch_outputs)
    total_pendientes = total_documentos - total_listos
    metrica_cargados, metrica_listos, metrica_pendientes, metrica_reutilizados = st.columns(4)
    metrica_cargados.metric("PDF cargados", total_documentos)
    metrica_listos.metric("Vistas listas", total_listos)
    metrica_pendientes.metric("Pendientes", total_pendientes)
    metrica_reutilizados.metric(
        "Reutilizados últimamente",
        st.session_state.last_cache_stats.get("reutilizados", 0)
    )
    st.caption(
        "Las vistas listas quedan guardadas durante la sesión. Cambiar de persona no ejecuta "
        "nuevamente la extracción ni la validación con IA."
    )
    col_generar_individual, col_generar_lote = st.columns(2)
    with col_generar_individual:
        if st.button(
            "👁️ Generar vista individual",
            disabled=bool(doc_actual.get("recomendaciones_pendientes_revision")),
            key=f"generate_preview_{id_editor}", use_container_width=True
        ):
            if not doc_actual.get("nombre") or not doc_actual.get("examenes_lista"):
                st.error("Completa como mínimo el nombre y los exámenes realizados.")
            else:
                try:
                    with st.spinner("Aplicando plantilla, firma y generando la vista final..."):
                        archivo_individual, huella_individual, reutilizado = obtener_o_generar_salida(
                            archivo_seleccionado, doc_actual, plantilla_activa, firma_activa,
                            formato_salida, st.session_state.batch_outputs,
                            st.session_state.output_fingerprints
                        )
                        st.session_state.batch_outputs[archivo_seleccionado] = archivo_individual
                        st.session_state.output_fingerprints[archivo_seleccionado] = huella_individual
                        st.session_state.processed_doc = archivo_individual
                        st.session_state.zip_bytes = construir_zip_salidas(st.session_state.batch_outputs)
                        st.session_state.last_cache_stats = {
                            "generados": 0 if reutilizado else 1,
                            "reutilizados": 1 if reutilizado else 0
                        }
                    st.success(
                        "Vista recuperada sin reprocesar." if reutilizado
                        else "Vista individual generada y guardada en el lote."
                    )
                except (RuntimeError, OSError, ValueError) as exc:
                    st.error(f"No se pudo preparar el documento: {exc}")
    with col_generar_lote:
        if st.button(
            "📚 Generar vista previa de todos los documentos",
            disabled=bool(incompletos_lote), key="generate_full_package_from_data",
            use_container_width=True
        ):
            try:
                progreso_lote = st.progress(0, text="Generando todos los documentos...")
                def actualizar_progreso_lote(
                    indice, total, nombre_pdf, datos_trabajador, reutilizado
                ):
                    accion = "Reutilizado" if reutilizado else "Generado"
                    progreso_lote.progress(
                        indice / total,
                        text=(
                            f"{accion} {indice} de {total}: "
                            f"{datos_trabajador.get('nombre') or nombre_pdf}"
                        )
                    )
                resultados_lote, zip_lote, huellas_lote, generados, reutilizados = preparar_lote_documentos(
                    st.session_state.documentos, plantilla_activa, firma_activa,
                    formato_salida, actualizar_progreso_lote,
                    st.session_state.batch_outputs, st.session_state.output_fingerprints
                )
                progreso_lote.empty()
                st.session_state.batch_outputs = resultados_lote
                st.session_state.output_fingerprints = huellas_lote
                st.session_state.zip_bytes = zip_lote
                st.session_state.last_cache_stats = {
                    "generados": generados, "reutilizados": reutilizados
                }
                if archivo_seleccionado in resultados_lote:
                    st.session_state.processed_doc = resultados_lote[archivo_seleccionado]
                st.success(
                    f"Lote listo: {generados} generado(s) y {reutilizados} reutilizado(s). "
                    "Los documentos están disponibles para descarga y correo."
                )
            except (RuntimeError, OSError, ValueError) as exc:
                st.error(
                    f"No fue posible completar el lote: {exc}. "
                    "Las vistas que ya estaban listas se conservaron."
                )

    if incompletos_lote:
        st.warning(
            "Para generar todo el lote, completa los datos y resuelve los fragmentos de: "
            + ", ".join(incompletos_lote)
        )

    st.caption(
        "Consulta el estado detallado de todo el paquete en ‘Control y tablas’."
    )

    if st.session_state.batch_outputs:
        st.markdown("### Vista previa y descargas")
        st.success(
            f"{len(st.session_state.batch_outputs)} de {len(st.session_state.documentos)} "
            "documentos están listos y permanecen disponibles al cambiar de persona."
        )
        documento_masivo_seleccionado = st.selectbox(
            "Selecciona el documento que deseas revisar",
            list(st.session_state.batch_outputs.keys()),
            format_func=lambda clave: st.session_state.documentos[clave].get("nombre") or clave,
            key="batch_preview_selector_data"
        )
        vista_masiva = st.session_state.batch_outputs[documento_masivo_seleccionado]
        if vista_masiva.get("preview_pdf"):
            try:
                mostrar_previsualizacion_pdf(vista_masiva["preview_pdf"])
            except Exception:
                components.html(vista_masiva["preview_html"], height=800, scrolling=True)
        else:
            components.html(vista_masiva["preview_html"], height=800, scrolling=True)
        descarga_individual, descarga_colectiva = st.columns(2)
        with descarga_individual:
            st.download_button(
                f"📄 Descargar {vista_masiva['filename']}",
                data=vista_masiva["bytes"], file_name=vista_masiva["filename"],
                mime=vista_masiva["mime"], key="download_selected_package_file",
                use_container_width=True
            )
        with descarga_colectiva:
            st.download_button(
                f"📦 Descargar {len(st.session_state.batch_outputs)} documentos en ZIP",
                data=st.session_state.zip_bytes,
                file_name=f"Lote_SST_JER_SA_{fecha.strftime('%Y%m%d')}.zip",
                mime="application/zip", key="download_full_package_zip_data",
                use_container_width=True
            )

with tab_origen:
    st.markdown("### Vista previa masiva de PDF originales")
    st.caption("Crea un índice de todos los PDF cargados y revísalos uno por uno, página por página, sin salir de la aplicación.")
    col_prev_masiva, col_cerrar_masiva = st.columns([3, 1])
    with col_prev_masiva:
        if st.button("👁️ Crear vista previa de todos los PDF cargados", key="create_original_batch_preview"):
            indice_originales = {}
            progreso_originales = st.progress(0, text="Preparando vista previa de los PDF originales...")
            total_originales = len(st.session_state.pdfs_raw_bytes)
            for numero, (nombre_original, contenido_original) in enumerate(st.session_state.pdfs_raw_bytes.items(), start=1):
                try:
                    with pdfplumber.open(io.BytesIO(contenido_original)) as documento_original:
                        indice_originales[nombre_original] = {
                            "paginas": len(documento_original.pages), "error": ""
                        }
                except Exception as exc:
                    indice_originales[nombre_original] = {"paginas": 0, "error": str(exc)}
                progreso_originales.progress(
                    numero / total_originales,
                    text=f"Indexado {numero} de {total_originales}: {nombre_original}"
                )
            progreso_originales.empty()
            st.session_state.original_batch_preview = indice_originales
    with col_cerrar_masiva:
        if st.session_state.original_batch_preview and st.button(
            "Cerrar vista masiva", key="close_original_batch_preview"
        ):
            st.session_state.original_batch_preview = {}

    if st.session_state.original_batch_preview:
        st.success(f"Vista masiva lista: {len(st.session_state.original_batch_preview)} PDF indexados.")
        st.caption("El índice completo está disponible en ‘Control y tablas’.")
        pdf_masivo_seleccionado = st.selectbox(
            "Selecciona un PDF del lote para revisarlo",
            list(st.session_state.original_batch_preview.keys()),
            format_func=lambda clave: (
                f"{st.session_state.documentos.get(clave, {}).get('nombre') or clave} · {clave}"
            ),
            key="original_batch_pdf_selector"
        )
        detalle_masivo = st.session_state.original_batch_preview[pdf_masivo_seleccionado]
        if detalle_masivo.get("error"):
            st.warning(f"No fue posible abrir este PDF: {detalle_masivo['error']}")
        elif detalle_masivo.get("paginas", 0):
            pagina_masiva = st.number_input(
                "Página que deseas visualizar", min_value=1,
                max_value=detalle_masivo["paginas"], value=1, step=1,
                key=f"original_batch_page_{hashlib.sha1(pdf_masivo_seleccionado.encode('utf-8')).hexdigest()[:8]}"
            )
            try:
                with pdfplumber.open(io.BytesIO(st.session_state.pdfs_raw_bytes[pdf_masivo_seleccionado])) as pdf_masivo:
                    pagina = pdf_masivo.pages[int(pagina_masiva) - 1]
                    st.image(
                        pagina.to_image(resolution=135).original,
                        caption=f"{pdf_masivo_seleccionado} · Página {int(pagina_masiva)} de {detalle_masivo['paginas']}",
                        use_container_width=True
                    )
            except Exception as exc:
                st.warning(f"No fue posible renderizar esta página: {exc}")

    st.markdown("---")
    st.markdown("### PDF original seleccionado")
    bytes_originales = st.session_state.pdfs_raw_bytes[archivo_seleccionado]
    st.download_button(
        "📥 Descargar PDF original", data=bytes_originales,
        file_name=f"ORIGINAL_{archivo_seleccionado}", mime="application/pdf", key=f"download_original_{id_editor}"
    )
    try:
        with pdfplumber.open(io.BytesIO(bytes_originales)) as preview_pdf:
            for i, page in enumerate(preview_pdf.pages):
                st.image(page.to_image(resolution=130).original, caption=f"Original · Página {i + 1}", use_container_width=True)
    except Exception as exc:
        st.warning(f"El PDF se cargó, pero no fue posible renderizarlo en pantalla: {exc}")

with tab_vista:
    doc_info = (
        st.session_state.batch_outputs.get(archivo_seleccionado)
        or st.session_state.processed_doc
    )
    if not doc_info:
        st.info("Primero revisa los datos y pulsa ‘Generar vista previa’. La descarga permanecerá bloqueada hasta entonces.")
    else:
        st.success(f"Documento listo · Consecutivo {doc_info['consec_num']}")
        if "HTML" in doc_info.get("format", ""):
            components.html(doc_info["preview_html"], height=850, scrolling=True)
        elif doc_info.get("preview_pdf"):
            try:
                mostrar_previsualizacion_pdf(doc_info["preview_pdf"])
            except Exception as exc:
                st.warning(f"No fue posible mostrar las páginas renderizadas: {exc}")
                components.html(doc_info["preview_html"], height=850, scrolling=True)
        else:
            st.warning("No fue posible convertir la plantilla a PDF para la vista fiel; se muestra una representación web.")
            components.html(doc_info["preview_html"], height=850, scrolling=True)
        st.download_button(
            label=f"📥 Descargar {doc_info['filename']}", data=doc_info["bytes"],
            file_name=doc_info["filename"], mime=doc_info["mime"], key=f"download_final_{id_editor}"
        )

with tab_lote:
    st.markdown("### Confirmación y envío de correos")
    st.caption(
        "Aquí seleccionas los destinatarios, personalizas el mensaje y confirmas el envío. "
        "El historial completo está en ‘Control y tablas’."
    )

    feedback_correo = st.session_state.get("email_send_feedback")
    if feedback_correo:
        if feedback_correo.get("exitos"):
            st.success(
                f"Último envío: {feedback_correo['exitos']} de {feedback_correo['total']} correo(s) enviados correctamente."
            )
        if feedback_correo.get("errores"):
            st.error("Errores del último envío:\n- " + "\n- ".join(feedback_correo["errores"]))

    if not st.session_state.batch_outputs:
        st.info(
            "Todavía no hay documentos preparados. En ‘Datos extraídos’ pulsa "
            "‘Generar vista previa de todos los documentos’."
        )
    else:
        st.success(f"Hay {len(st.session_state.batch_outputs)} documento(s) listos para enviar.")
        st.markdown("#### 1. Selecciona el alcance")
        modo_envio = st.radio(
            "Alcance del envío", ["Un correo", "Varios correos"],
            horizontal=True, key="email_scope"
        )
        if modo_envio == "Un correo":
            documento_correo = st.selectbox(
                "Documento que deseas enviar",
                list(st.session_state.batch_outputs.keys()),
                format_func=lambda clave: st.session_state.documentos[clave].get("nombre") or clave,
                key="single_email_document"
            )
            seleccion_envio = [documento_correo]
        else:
            seleccion_envio = st.multiselect(
                "Documentos que deseas enviar",
                list(st.session_state.batch_outputs.keys()),
                default=list(st.session_state.batch_outputs.keys()),
                format_func=lambda clave: st.session_state.documentos[clave].get("nombre") or clave,
                key="email_document_selection"
            )

        st.markdown("#### 2. Configura copias y mensaje")
        columna_cc, columna_cco = st.columns(2)
        with columna_cc:
            cc_entrada = st.text_input(
                "Con copia (CC)", key="email_cc_global",
                placeholder="correo1@empresa.com; correo2@empresa.com",
                help="Estos destinatarios serán visibles en el correo."
            )
        with columna_cco:
            cco_entrada = st.text_input(
                "Con copia oculta (CCO)", key="email_bcc_global",
                placeholder="auditoria@empresa.com",
                help="Estos destinatarios recibirán el correo sin ser visibles para los demás."
            )
        st.caption(
            "En un envío masivo, las copias configuradas se incluirán en cada correo. "
            "Puedes separar varias direcciones con coma, punto y coma o salto de línea."
        )
        asunto_plantilla = st.text_input(
            "Asunto del correo", value=PLANTILLA_ASUNTO_CORREO,
            key="email_subject_template",
            help="Puedes usar {nombre} y {identificacion}; se reemplazan para cada persona."
        )
        cuerpo_plantilla = st.text_area(
            "Cuerpo del mensaje", value=PLANTILLA_CUERPO_CORREO,
            height=270, key="email_body_template",
            help="Edita libremente el mensaje. Conserva {nombre} y {identificacion} si quieres personalización automática."
        )

        envios_preparados = []
        resumen_confirmacion = []
        for clave_documento in seleccion_envio:
            datos_envio = st.session_state.documentos[clave_documento]
            archivo_envio = st.session_state.batch_outputs[clave_documento]
            widget_id = hashlib.sha1(clave_documento.encode("utf-8")).hexdigest()[:8]
            destinatario = st.text_input(
                f"Correo de {datos_envio.get('nombre') or clave_documento}",
                value=datos_envio.get("correo", ""), key=f"email_to_{widget_id}"
            ).strip().lower()
            datos_envio["correo"] = destinatario
            asunto = generar_asunto_correo(datos_envio, asunto_plantilla).strip()
            mensaje = generar_mensaje_correo(datos_envio, cuerpo_plantilla).strip()
            envio_preparado = {
                "clave": clave_documento,
                "destinatario": destinatario,
                "cc": separar_correos(cc_entrada),
                "cco": separar_correos(cco_entrada),
                "asunto": asunto,
                "mensaje": mensaje,
                "archivo": archivo_envio
            }
            envios_preparados.append(envio_preparado)
            resumen_confirmacion.append({
                "Trabajador": datos_envio.get("nombre", ""),
                "Destinatario": destinatario,
                "CC": ", ".join(envio_preparado["cc"]) or "—",
                "CCO": ", ".join(envio_preparado["cco"]) or "—",
                "Asunto": asunto,
                "Adjunto": archivo_envio["filename"],
                "Estado": "Listo" if correo_valido(destinatario) else "Revisar correo"
            })

        if resumen_confirmacion:
            st.markdown("#### 3. Revisa la confirmación")
            for indice_envio, envio in enumerate(envios_preparados):
                resumen = resumen_confirmacion[indice_envio]
                estado_color = "#4ade80" if resumen["Estado"] == "Listo" else "#fbbf24"
                st.markdown(
                    "<div class='email-card'>"
                    f"<strong>{html.escape(resumen['Trabajador'] or 'Sin nombre')}</strong> "
                    f"<span style='color:{estado_color}'>● {html.escape(resumen['Estado'])}</span><br>"
                    f"<small>Para: {html.escape(resumen['Destinatario'] or 'Sin correo')} · "
                    f"Adjunto: {html.escape(resumen['Adjunto'])}</small>"
                    "</div>",
                    unsafe_allow_html=True
                )
                with st.expander(
                    f"Ver mensaje de {resumen['Trabajador'] or 'este destinatario'}",
                    expanded=(len(envios_preparados) == 1)
                ):
                    st.markdown(f"**Asunto:** {envio['asunto']}")
                    st.text(envio["mensaje"])
                    st.caption(
                        f"CC: {resumen['CC']} · CCO: {resumen['CCO']}"
                    )

        smtp_listo = bool(
            configuracion_smtp.get("host") and configuracion_smtp.get("sender")
            and configuracion_smtp.get("user") and configuracion_smtp.get("password")
        )
        if not smtp_listo:
            st.warning("Completa servidor, usuario, remitente y clave de aplicación en ‘Correo saliente’.")
        destinatarios_invalidos = [
            st.session_state.documentos[envio["clave"]].get("nombre") or envio["clave"]
            for envio in envios_preparados if not correo_valido(envio["destinatario"])
        ]
        copias_invalidas = correos_invalidos(cc_entrada) + correos_invalidos(cco_entrada)
        if destinatarios_invalidos:
            st.warning("Falta revisar el correo de: " + ", ".join(destinatarios_invalidos))
        if copias_invalidas:
            st.warning("Revisa los correos de CC o CCO: " + ", ".join(copias_invalidas))
        mensaje_incompleto = not asunto_plantilla.strip() or not cuerpo_plantilla.strip()
        if mensaje_incompleto:
            st.warning("El asunto y el cuerpo del mensaje no pueden quedar vacíos.")

        with st.form("email_send_confirmation_form", clear_on_submit=False):
            confirmar_envio = st.checkbox(
                f"Confirmo el envío de {len(envios_preparados)} correo(s), con sus copias y documentos correspondientes.",
                value=False, key="confirm_email_send_form"
            )
            enviar_confirmado = st.form_submit_button(
                "📨 Confirmar y enviar correos",
                disabled=not (
                    envios_preparados and smtp_listo and not destinatarios_invalidos
                    and not copias_invalidas and not mensaje_incompleto
                ),
                use_container_width=True
            )

        if enviar_confirmado:
            if not confirmar_envio:
                st.warning("Marca la casilla de confirmación antes de enviar.")
            else:
                exitos = 0
                errores = []
                progreso_correo = st.progress(0, text="Enviando correos...")
                for indice, envio in enumerate(envios_preparados, start=1):
                    clave = envio["clave"]
                    destinatario = envio["destinatario"]
                    asunto = envio["asunto"]
                    mensaje = envio["mensaje"]
                    archivo = envio["archivo"]
                    trabajador_envio = st.session_state.documentos[clave].get("nombre") or clave
                    detalle_envio = ""
                    try:
                        enviar_archivo_por_correo(
                            destinatario, asunto, mensaje, archivo, configuracion_smtp,
                            cc=envio["cc"], cco=envio["cco"]
                        )
                        exitos += 1
                        estado = "Enviado"
                        detalle_envio = "Mensaje aceptado por el servidor SMTP."
                    except (smtplib.SMTPException, ssl.SSLError, OSError, ValueError) as exc:
                        estado = "Error"
                        detalle_envio = str(exc)
                        errores.append(f"{trabajador_envio}: {exc}")
                    registrar_historial_correo(
                        clave, trabajador_envio, destinatario, archivo["filename"],
                        estado, detalle_envio, cc=", ".join(envio["cc"]),
                        cco=", ".join(envio["cco"]), asunto=asunto
                    )
                    progreso_correo.progress(
                        indice / len(envios_preparados),
                        text=f"Procesado {indice} de {len(envios_preparados)}: {trabajador_envio}"
                    )
                progreso_correo.empty()
                st.session_state.email_send_feedback = {
                    "exitos": exitos, "total": len(envios_preparados), "errores": errores
                }
                if exitos:
                    st.success(f"Se enviaron correctamente {exitos} de {len(envios_preparados)} correos.")
                if errores:
                    st.error("No se enviaron algunos correos:\n- " + "\n- ".join(errores))

with tab_control:
    st.markdown("### Centro de control y tablas")
    st.caption(
        "Las tablas técnicas y de seguimiento están reunidas aquí para que el flujo "
        "de revisión, generación y correo sea más limpio."
    )
    control_validacion, control_paquete, control_originales, control_correos = st.tabs([
        "Validación", "Paquete final", "PDF originales", "Correos"
    ])

    with control_validacion:
        st.dataframe([
            {
                "PDF": nombre_pdf,
                "Motor": datos.get("modo_validacion", "Respaldo local"),
                "Exámenes": len(datos.get("examenes_lista", [])),
                "Recomendaciones": len(datos.get("recomendaciones_lista", [])),
                "Fragmentos pendientes": len(datos.get("recomendaciones_pendientes_revision", [])),
                "Versión": PROCESSING_PIPELINE_VERSION
            }
            for nombre_pdf, datos in st.session_state.documentos.items()
        ], use_container_width=True, hide_index=True)

    with control_paquete:
        st.dataframe([
            {
                "Trabajador": datos.get("nombre") or "Sin nombre",
                "PDF origen": nombre_pdf,
                "Estado": "Listo" if nombre_pdf in st.session_state.batch_outputs else "Pendiente",
                "Archivo final": st.session_state.batch_outputs.get(nombre_pdf, {}).get("filename", "—")
            }
            for nombre_pdf, datos in st.session_state.documentos.items()
        ], use_container_width=True, hide_index=True)

    with control_originales:
        if st.session_state.original_batch_preview:
            st.dataframe([
                {
                    "PDF": nombre_original,
                    "Trabajador": st.session_state.documentos.get(nombre_original, {}).get("nombre", ""),
                    "Páginas": detalle.get("paginas", 0),
                    "Estado": "Listo" if not detalle.get("error") else "No renderizable"
                }
                for nombre_original, detalle in st.session_state.original_batch_preview.items()
            ], use_container_width=True, hide_index=True)
        else:
            st.info("Crea primero el índice desde ‘PDF original’ para consultar esta tabla.")

    with control_correos:
        if resumen_confirmacion:
            st.markdown("#### Preparación actual")
            st.dataframe(resumen_confirmacion, use_container_width=True, hide_index=True)
        st.markdown("#### Historial de envíos")
        historial_persistente = obtener_historial_correos()
        if historial_persistente:
            st.dataframe(historial_persistente, use_container_width=True, hide_index=True)
        else:
            st.info("Todavía no se han registrado intentos de envío.")
